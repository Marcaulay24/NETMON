#!/usr/bin/env python3
"""
🛡️ NETMON PRO - Enterprise Network Security Suite v2.1
✅ FULL DEVICE DETAILS | MAC | HOSTNAME | VENDOR LOOKUP | DETAILS PANEL
✅ SQLITE PERSISTENCE | OS FINGERPRINTING | VULN MAPPING
✅ Authorized Pentest Tool - All Features Unlocked
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext, simpledialog
import threading
import json
import time
from datetime import datetime
import psutil
import os
import socket
import subprocess
import ipaddress
import queue
import csv
import re
import sqlite3
import platform
import hashlib
import uuid
from collections import defaultdict

# Optional dependencies
SCAPY_AVAILABLE = False
NMAP_AVAILABLE = False
FPDF_AVAILABLE = False
DOCX_AVAILABLE = False
nmap = None
docx = None
sniff = None
ARP = None
IP = None
DNSQR = None
TCP = None
Raw = None

def _resolve_fpdf_class():
    try:
        from fpdf import FPDF
        return FPDF
    except ImportError:
        return None

def check_dependencies():
    global SCAPY_AVAILABLE, NMAP_AVAILABLE, FPDF_AVAILABLE, DOCX_AVAILABLE
    global nmap, docx, sniff, ARP, IP, DNSQR, TCP, Raw
    SCAPY_AVAILABLE = False
    NMAP_AVAILABLE = False
    DOCX_AVAILABLE = False
    nmap = None
    docx = None
    sniff = None
    ARP = None
    IP = None
    DNSQR = None
    TCP = None
    Raw = None
    try:
        from scapy.all import sniff as scapy_sniff, ARP as scapy_ARP, IP as scapy_IP
        from scapy.all import DNSQR as scapy_DNSQR, TCP as scapy_TCP, Raw as scapy_Raw
        sniff = scapy_sniff
        ARP = scapy_ARP
        IP = scapy_IP
        DNSQR = scapy_DNSQR
        TCP = scapy_TCP
        Raw = scapy_Raw
        SCAPY_AVAILABLE = True
    except Exception:
        SCAPY_AVAILABLE = False
    try:
        import nmap as nmap_module
        nmap = nmap_module
        NMAP_AVAILABLE = True
    except ImportError: pass
    
    # fpdf2 installs as import name "fpdf"
    FPDF_AVAILABLE = _resolve_fpdf_class() is not None

    try:
        import docx as docx_module
        docx = docx_module
        DOCX_AVAILABLE = True
    except ImportError: pass

check_dependencies()

DB_FILE = "netmon_history.db"

class AlertManager:
    def __init__(self, on_alert=None):
        self.on_alert = on_alert
        self.alerts = []
        self.alert_levels = {
            "critical": {"rule_range": (15, 20), "color": "#c53030", "sound": True},
            "high": {"rule_range": (12, 14), "color": "#d69e2e", "sound": True},
            "medium": {"rule_range": (7, 11), "color": "#4299e1", "sound": False},
            "low": {"rule_range": (0, 6), "color": "#0f766e", "sound": False},
        }

    def categorize_severity(self, rule_level):
        for severity, config in self.alert_levels.items():
            low, high = config["rule_range"]
            if low <= rule_level <= high:
                return severity
        return "low"

    def generate_correlation_id(self):
        return uuid.uuid4().hex[:12]

    def process_alert(self, rule_level, message, source_ip):
        severity = self.categorize_severity(int(rule_level))
        alert = {
            "timestamp": datetime.now(),
            "severity": severity,
            "rule_level": int(rule_level),
            "message": message,
            "source_ip": source_ip,
            "status": "new",
            "correlation_id": self.generate_correlation_id(),
        }
        self.alerts.append(alert)
        if len(self.alerts) > 5000:
            self.alerts = self.alerts[-5000:]

        if severity == "critical":
            self.send_slack_alert(alert)
            self.trigger_siem_forward(alert)
            self.create_ticket(alert)
            self.execute_playbook(alert)

        if self.on_alert:
            self.on_alert(alert)
        return alert

    def send_slack_alert(self, alert):
        return False

    def trigger_siem_forward(self, alert):
        return False

    def create_ticket(self, alert):
        return False

    def execute_playbook(self, alert):
        return False

    def counts_last_24h(self):
        now = datetime.now()
        counts = {"critical": 0, "high": 0, "medium": 0, "low": 0}
        for alert in self.alerts:
            if (now - alert["timestamp"]).total_seconds() <= 86400:
                counts[alert["severity"]] = counts.get(alert["severity"], 0) + 1
        return counts


class FileMonitor:
    def __init__(self):
        self.baseline = {}
        self.monitored_paths = []

    def _file_digest(self, path):
        h = hashlib.sha256()
        with open(path, "rb") as f:
            for chunk in iter(lambda: f.read(8192), b""):
                h.update(chunk)
        return h.hexdigest()

    def _file_meta(self, path):
        st = os.stat(path)
        return {
            "hash": self._file_digest(path),
            "mode": oct(st.st_mode & 0o777),
            "uid": st.st_uid,
            "gid": st.st_gid,
            "size": st.st_size,
            "mtime": int(st.st_mtime),
        }

    def baseline_filesystem(self, paths):
        baseline = {}
        for root in paths:
            if not os.path.exists(root):
                continue
            for dirpath, _, files in os.walk(root):
                for name in files:
                    full = os.path.join(dirpath, name)
                    try:
                        baseline[full] = self._file_meta(full)
                    except Exception:
                        continue
        self.baseline = baseline
        self.monitored_paths = list(paths)
        return baseline

    def fim_alert_rules(self, event_type, path):
        critical_keywords = ("passwd", "shadow", "authorized_keys", "sudoers", "hosts")
        is_critical = any(k in path.lower() for k in critical_keywords)
        if event_type in ("deleted", "created") and is_critical:
            return "critical", 15
        if event_type in ("content_changed", "permission_changed", "owner_changed"):
            return ("high", 13) if is_critical else ("medium", 9)
        return "low", 4

    def detect_file_changes(self):
        current = {}
        changes = []
        for root in self.monitored_paths:
            if not os.path.exists(root):
                continue
            for dirpath, _, files in os.walk(root):
                for name in files:
                    full = os.path.join(dirpath, name)
                    try:
                        current[full] = self._file_meta(full)
                    except Exception:
                        continue

        old_paths = set(self.baseline.keys())
        new_paths = set(current.keys())

        for created in sorted(new_paths - old_paths):
            severity, rule_level = self.fim_alert_rules("created", created)
            changes.append({"type": "created", "path": created, "severity": severity, "rule_level": rule_level})
        for deleted in sorted(old_paths - new_paths):
            severity, rule_level = self.fim_alert_rules("deleted", deleted)
            changes.append({"type": "deleted", "path": deleted, "severity": severity, "rule_level": rule_level})

        for common in sorted(old_paths & new_paths):
            old = self.baseline[common]
            new = current[common]
            if old["hash"] != new["hash"]:
                severity, rule_level = self.fim_alert_rules("content_changed", common)
                changes.append({"type": "content_changed", "path": common, "severity": severity, "rule_level": rule_level})
            if old["mode"] != new["mode"]:
                severity, rule_level = self.fim_alert_rules("permission_changed", common)
                changes.append({"type": "permission_changed", "path": common, "severity": severity, "rule_level": rule_level})
            if old["uid"] != new["uid"] or old["gid"] != new["gid"]:
                severity, rule_level = self.fim_alert_rules("owner_changed", common)
                changes.append({"type": "owner_changed", "path": common, "severity": severity, "rule_level": rule_level})

        self.baseline = current
        return changes


class ThreatHunter:
    def __init__(self):
        self.mitre_framework = {
            "TA0001": "Initial Access",
            "TA0002": "Execution",
            "TA0003": "Persistence",
            "TA0004": "Privilege Escalation",
            "TA0005": "Defense Evasion",
            "TA0006": "Credential Access",
            "TA0007": "Discovery",
            "TA0008": "Lateral Movement",
            "TA0009": "Collection",
            "TA0010": "Exfiltration",
            "TA0011": "Command and Control",
        }

    def map_to_mitre(self, text):
        t = text.lower()
        if "brute" in t or "rdp" in t:
            return "TA0006", "Credential Access"
        if "arp" in t or "mitm" in t:
            return "TA0008", "Lateral Movement"
        if "dns" in t or "http" in t:
            return "TA0011", "Command and Control"
        if "malware" in t:
            return "TA0002", "Execution"
        return "TA0007", "Discovery"

    def threat_score_device(self, details, info):
        score = 0
        ports = set(int(p) for p in details.get("ports", {}).keys())
        if ports & {21, 23, 445, 3389}:
            score += 35
        if details.get("os", "").lower().startswith("unknown"):
            score += 15
        if "LIVE" in info.get("status", ""):
            score += 10
        if info.get("is_critical", False):
            score += 25
        return min(score, 100)


class DockerMonitor:
    def monitor_containers(self):
        return {"status": "unavailable", "reason": "docker integration optional"}


class AWSMonitor:
    def monitor_aws_events(self):
        return {"status": "unavailable", "reason": "aws integration optional"}


class GitHubMonitor:
    def monitor_audit_log(self):
        return {"status": "unavailable", "reason": "github integration optional"}


class ComplianceReporter:
    def __init__(self, app):
        self.app = app

    def generate_framework_report(self, framework="pci"):
        framework = framework.lower()
        templates = self.app.compliance_templates
        if framework in ("pci", "pci_dss"):
            controls = templates.get("pci_dss", [])
        elif framework == "hipaa":
            controls = templates.get("hipaa", [])
        elif framework == "gdpr":
            controls = templates.get("gdpr", [])
        else:
            controls = templates.get("pci_dss", []) + templates.get("hipaa", []) + templates.get("gdpr", [])
        return {
            "framework": framework,
            "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "controls": controls,
            "grc_score": self.app.grc_results.get("score", "N/A"),
            "grc_risk": self.app.grc_results.get("risk_level", "N/A"),
            "hygiene_score": self.app.hygiene_results.get("overall", "N/A"),
        }

    def generate_audit_trail(self):
        return {
            "alerts_24h": self.app.alert_manager.counts_last_24h(),
            "config_assessments": len(self.app.config_assessment_results),
            "grc_assessment": self.app.grc_results.get("timestamp", "N/A"),
            "security_audit": self.app.last_audit_summary,
        }

class NetMonPro:
    def __init__(self, root):
        self.root = root
        self.root.title("🛡️ NetMon PRO v2.1 - Enterprise Pentest Suite")
        self.root.geometry("1800x1100")
        self.root.minsize(1400, 900)
        self.root.configure(bg='#1e1e1e')
        
        # Data storage
        self.devices = {}
        self.device_details = {}
        self.router_ip = None
        self.isp_range = []
        self.scan_results = {}
        self.traffic_logs = {} # ip -> [dns_queries]
        self.grc_results = {}
        self.last_audit_summary = {}
        self.config_assessment_results = {}
        self.hygiene_results = {}
        self.fim_running = False
        self.monitoring_paused = False
        self.log_queue = queue.Queue()
        self.ui_queue = queue.Queue()
        self.selected_device = None
        self.main_thread_id = threading.get_ident()
        self.device_filter_var = tk.StringVar(value="")
        self.live_only_var = tk.BooleanVar(value=False)
        self.state_lock = threading.RLock()
        
        # UI Adjustment Variables
        self.ui_width = tk.IntVar(value=1600)
        self.ui_height = tk.IntVar(value=900)
        self.left_panel_width = tk.IntVar(value=350)
        self.right_panel_width = tk.IntVar(value=550)
        self.fim_interval = tk.IntVar(value=20)
        self.auto_remediate = tk.BooleanVar(value=False)
        self.default_fim_path = tk.StringVar(value="/etc" if platform.system() != "Windows" else "C:\\Windows\\System32")

        check_dependencies()
        self.alert_manager = AlertManager(on_alert=self._handle_new_alert)
        self.file_monitor = FileMonitor()
        self.threat_hunter = ThreatHunter()
        self.cloud_monitors = {
            "docker": DockerMonitor(),
            "aws": AWSMonitor(),
            "github": GitHubMonitor(),
        }
        self.compliance_reporter = ComplianceReporter(self)
        self.compliance_templates = {
            "pci_dss": ["Requirement 1: Firewall", "Requirement 2: Secure Config", "Requirement 10: Logging"],
            "hipaa": ["Access Control", "Audit Controls", "Integrity", "Transmission Security"],
            "gdpr": ["Data Minimization", "Lawful Processing", "Breach Notification"],
        }
        
        # Initialize Database
        self._init_db()
        
        self.root.after(100, self._process_log_queue)
        self.root.after(100, self._process_ui_queue)
        self._configure_styles()
        self._create_gui()
        self._start_monitoring()
        self._start_traffic_monitor()
        self._auto_discover_network()
        self._load_history()
        
        self.log("🚀 NetMon PRO v2.1 - Enterprise Pentest Ready!")
        self.log("📦 SQLite Persistence: Enabled")
    
    def _init_db(self):
        """Initialize SQLite database for device history"""
        try:
            with sqlite3.connect(DB_FILE) as conn:
                c = conn.cursor()
                c.execute('''CREATE TABLE IF NOT EXISTS devices
                             (ip TEXT PRIMARY KEY, mac TEXT, hostname TEXT, 
                              vendor TEXT, status TEXT, last_seen TEXT,
                              custom_name TEXT, connection_time TEXT, is_critical INTEGER DEFAULT 0)''')

                # Migration: Add columns if they don't exist
                try:
                    c.execute("ALTER TABLE devices ADD COLUMN custom_name TEXT")
                except sqlite3.OperationalError:
                    pass
                try:
                    c.execute("ALTER TABLE devices ADD COLUMN connection_time TEXT")
                except sqlite3.OperationalError:
                    pass
                try:
                    c.execute("ALTER TABLE devices ADD COLUMN is_critical INTEGER DEFAULT 0")
                except sqlite3.OperationalError:
                    pass
        except Exception as e:
            self.log(f"❌ DB Init Error: {e}", "ERROR")

    def _save_to_db(self, ip, info):
        """Save device data to history"""
        try:
            with sqlite3.connect(DB_FILE) as conn:
                c = conn.cursor()
                c.execute('''INSERT OR REPLACE INTO devices VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)''',
                          (ip, info.get('mac'), info.get('hostname'), 
                           info.get('vendor'), info.get('status'), 
                           datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                           info.get('custom_name', ''), info.get('connection_time', ''),
                           1 if info.get('is_critical', False) else 0))
        except Exception as exc:
            self.log(f"⚠️ DB Save Error ({ip}): {exc}", "WARN")

    def _load_history(self):
        """Load previously discovered devices"""
        try:
            with sqlite3.connect(DB_FILE) as conn:
                c = conn.cursor()
                c.execute("SELECT * FROM devices")
                for row in c.fetchall():
                    ip, mac, host, vendor, status, last, custom, conn_time, critical = row
                    mac_value = self._normalize_mac(mac) or ('??:??:??:??:??:??' if self._is_unknown_identity(mac) else str(mac or '??:??:??:??:??:??'))
                    host_value = host if not self._is_unknown_identity(host) else 'Unknown'
                    vendor_value = vendor if not self._is_unknown_identity(vendor) else 'Unknown'
                    self.devices[ip] = {
                        'mac': mac_value, 'hostname': host_value,
                        'vendor': vendor_value, 'status': f'OFFLINE (Last: {last})',
                        'custom_name': custom or '',
                        'connection_time': conn_time or '',
                        'is_critical': bool(critical)
                    }
            self._update_devices_tree()
            self.log(f"📚 Loaded {len(self.devices)} devices from history")
        except Exception as exc:
            self.log(f"⚠️ History Load Error: {exc}", "WARN")

    def log(self, message, level="INFO"):
        timestamp = datetime.now().strftime("[%H:%M:%S]")
        icon = ""
        if level == "SUCCESS": icon = "✓ "
        elif level == "ERROR" or level == "CRITICAL": icon = "✖ "
        elif level == "WARN": icon = "⚠ "
        
        full_message = f"{timestamp} {icon}{message}\n"
        self.log_queue.put((full_message, level))
    
    def _process_log_queue(self):
        try:
            while True:
                item = self.log_queue.get_nowait()
                msg, level = item if isinstance(item, tuple) else (item, "INFO")
                
                if hasattr(self, "log_text") and self.log_text.winfo_exists():
                    # Insert timestamp with blue color
                    ts_end = msg.find("]") + 1
                    self.log_text.insert(tk.END, msg[:ts_end], 'TIMESTAMP')
                    self.log_text.insert(tk.END, msg[ts_end:], level)
                    self.log_text.see(tk.END) # Auto-scroll
        except queue.Empty:
            pass
        self.root.after(100, self._process_log_queue)

    def _process_ui_queue(self):
        try:
            while True:
                callback, args, kwargs = self.ui_queue.get_nowait()
                callback(*args, **kwargs)
        except queue.Empty:
            pass
        self.root.after(100, self._process_ui_queue)

    def _queue_ui(self, callback, *args, **kwargs):
        if threading.get_ident() == self.main_thread_id:
            callback(*args, **kwargs)
            return
        self.ui_queue.put((callback, args, kwargs))

    def _set_text_widget(self, widget, content):
        widget.delete(1.0, tk.END)
        widget.insert(tk.END, content)

    def _configure_styles(self):
        style = ttk.Style()
        if "clam" in style.theme_names():
            style.theme_use("clam")
        style.configure("TFrame", background="#161b22")
        style.configure("TLabelframe", background="#161b22", foreground="#58a6ff")
        style.configure("TLabelframe.Label", background="#161b22", foreground="#58a6ff", font=("Segoe UI", 10, "bold"))
        style.configure("TButton", font=("Segoe UI", 10, "bold"), padding=6, background="#21262d", foreground="#c9d1d9")
        style.map("TButton", 
                  background=[("active", "#30363d"), ("pressed", "#0d1117")], 
                  foreground=[("active", "#58a6ff"), ("pressed", "#58a6ff")])
        style.configure("TCheckbutton", background="#161b22", foreground="#c9d1d9")
        style.configure("TNotebook", background="#0d1117", borderwidth=0)
        style.configure("TNotebook.Tab", background="#21262d", foreground="#8b949e", font=("Segoe UI", 9, "bold"), padding=(8, 6))
        style.map("TNotebook.Tab", 
                  background=[("selected", "#58a6ff")], 
                  foreground=[("selected", "#0d1117")])
        style.configure("Treeview", background="#0d1117", fieldbackground="#0d1117", foreground="#c9d1d9", rowheight=28, borderwidth=0)
        style.configure("Treeview.Heading", background="#21262d", foreground="#58a6ff", font=("Segoe UI", 10, "bold"))

    @staticmethod
    def _is_valid_ipv4(value):
        try:
            ipaddress.IPv4Address(value)
            return True
        except ValueError:
            return False

    @staticmethod
    def _sanitize_host(value, max_len=120):
        cleaned = re.sub(r"[^a-zA-Z0-9.\-_]", "", value.strip())
        return cleaned[:max_len]

    @staticmethod
    def _sanitize_label(value, max_len=48):
        cleaned = re.sub(r"[^\w\s\-.:()/]", "", value.strip())
        return cleaned[:max_len]

    def _set_footer(self, text):
        self.footer_label.config(text=text)

    def _start_task_ui(self, label):
        self._set_footer(label)
        self.progress_bar.start(10)

    def _end_task_ui(self, label="Ready"):
        self.progress_bar.stop()
        self._set_footer(label)

    def _run_command(self, args, timeout=10, capture_output=False, silent=False):
        kwargs = {"timeout": timeout}
        if capture_output:
            kwargs.update({"capture_output": True, "text": True})
        else:
            kwargs.update({"stdout": subprocess.DEVNULL, "stderr": subprocess.DEVNULL})
        try:
            return subprocess.run(args, **kwargs)
        except FileNotFoundError:
            if not silent:
                self.log(f"⚠️ Required command not found: {args[0]}", "WARN")
            return None
        except subprocess.TimeoutExpired:
            if not silent:
                self.log(f"⚠️ Command timeout: {' '.join(args)}", "WARN")
            return None

    @staticmethod
    def _is_root_user():
        if hasattr(os, "getuid"):
            return os.getuid() == 0
        return False

    def _save_alert_to_db(self, alert):
        try:
            with sqlite3.connect(DB_FILE) as conn:
                conn.execute(
                    """CREATE TABLE IF NOT EXISTS alerts (
                        correlation_id TEXT PRIMARY KEY,
                        timestamp TEXT,
                        severity TEXT,
                        rule_level INTEGER,
                        message TEXT,
                        source_ip TEXT,
                        status TEXT
                    )"""
                )
                conn.execute(
                    """INSERT OR REPLACE INTO alerts
                       (correlation_id, timestamp, severity, rule_level, message, source_ip, status)
                       VALUES (?, ?, ?, ?, ?, ?, ?)""",
                    (
                        alert["correlation_id"],
                        alert["timestamp"].strftime("%Y-%m-%d %H:%M:%S"),
                        alert["severity"],
                        alert["rule_level"],
                        alert["message"],
                        alert["source_ip"],
                        alert["status"],
                    ),
                )
        except Exception as exc:
            self.log(f"⚠️ Failed to persist alert: {exc}", "WARN")

    def _handle_new_alert(self, alert):
        self._save_alert_to_db(alert)
        message = (
            f"[{alert['timestamp'].strftime('%H:%M:%S')}] "
            f"{alert['severity'].upper():<8} L{alert['rule_level']:<2} "
            f"{alert['source_ip']:<15} {alert['message']}"
        )
        self.log(f"🚨 Alert[{alert['severity']}]: {alert['message']} ({alert['source_ip']})")
        self._queue_ui(self._append_text_widget, self.alerts_text, message + "\n")
        self._queue_ui(self._update_overview_dashboard)

    def _emit_security_alert(self, rule_level, message, source_ip="system"):
        self.alert_manager.process_alert(rule_level=rule_level, message=message, source_ip=source_ip)

    def _append_text_widget(self, widget, content):
        widget.insert(tk.END, content)
        widget.see(tk.END)

    @staticmethod
    def _wheel_steps(event):
        if getattr(event, "num", None) == 4:
            return -1
        if getattr(event, "num", None) == 5:
            return 1
        delta = getattr(event, "delta", 0)
        if delta == 0:
            return 0
        if platform.system() == "Darwin":
            return -1 if delta > 0 else 1
        steps = int(-delta / 120)
        return steps if steps != 0 else (-1 if delta > 0 else 1)

    def _smooth_scroll_widget(self, widget, steps):
        direction = 1 if steps > 0 else -1
        # Increased pulses for faster response
        pulses = min(max(abs(steps) * 3, 2), 12)

        def smooth_scroll(remaining):
            if remaining <= 0:
                return
            try:
                widget.yview_scroll(direction, "units")
            except Exception:
                return
            widget.after(10, lambda: smooth_scroll(remaining - 1))

        smooth_scroll(pulses)

    def _enable_smooth_scroll(self, widget):

        def on_wheel(event):
            steps = self._wheel_steps(event)
            if steps != 0:
                self._smooth_scroll_widget(widget, steps)
            return "break"
        widget.bind("<MouseWheel>", on_wheel, add="+")
        widget.bind("<Button-4>", on_wheel, add="+")
        widget.bind("<Button-5>", on_wheel, add="+")

    def _bind_master_wheel_targets(self, parent):
        # Expanded target list for comprehensive master scroll coverage
        passive = (
            tk.Frame, tk.Label, tk.LabelFrame, tk.Canvas,
            ttk.Frame, ttk.Label, ttk.Labelframe, ttk.Button, ttk.Notebook,
            ttk.Entry, ttk.Checkbutton, ttk.Scrollbar
        )

        def on_master_wheel(event):
            steps = self._wheel_steps(event)
            if steps != 0:
                self._smooth_scroll_widget(self.app_canvas, steps)
            return "break"

        for child in parent.winfo_children():
            if isinstance(child, passive):
                child.bind("<MouseWheel>", on_master_wheel, add="+")
                child.bind("<Button-4>", on_master_wheel, add="+")
                child.bind("<Button-5>", on_master_wheel, add="+")
            self._bind_master_wheel_targets(child)

    def _setup_master_scroll(self):
        self.app_canvas = tk.Canvas(self.root, bg="#161b22", highlightthickness=0)
        self.v_scrollbar = ttk.Scrollbar(self.root, orient="vertical", command=self.app_canvas.yview)
        self.h_scrollbar = ttk.Scrollbar(self.root, orient="horizontal", command=self.app_canvas.xview)
        
        self.app_canvas.configure(yscrollcommand=self.v_scrollbar.set, xscrollcommand=self.h_scrollbar.set)
        
        self.v_scrollbar.pack(side="right", fill="y")
        self.h_scrollbar.pack(side="bottom", fill="x")
        self.app_canvas.pack(side="left", fill="both", expand=True)

        self.app_viewport = tk.Frame(self.app_canvas, bg="#161b22")
        self.app_window = self.app_canvas.create_window((0, 0), window=self.app_viewport, anchor="nw")

        def on_viewport_configure(_event):
            # Update the scrollregion to encompass the inner frame
            self.app_canvas.configure(scrollregion=self.app_canvas.bbox("all"))

        def on_canvas_configure(event):
            # Resize the inner frame to match the canvas width
            self.app_canvas.itemconfigure(self.app_window, width=event.width)

        self.app_viewport.bind("<Configure>", on_viewport_configure, add="+")
        self.app_canvas.bind("<Configure>", on_canvas_configure, add="+")
        self._enable_smooth_scroll(self.app_canvas)

    def _select_tab_by_title(self, title):
        try:
            target = re.sub(r"[^a-z0-9]+", "", title.lower())
            tabs = self.notebook.tabs()
            for tab_id in tabs:
                tab_text = self.notebook.tab(tab_id, "text")
                normalized = re.sub(r"[^a-z0-9]+", "", tab_text.lower())
                if tab_text == title or normalized == target or target in normalized:
                    self.notebook.select(tab_id)
                    return True
        except Exception:
            return False
        return False

    def _show_cloud_status(self, provider):
        monitor = self.cloud_monitors.get(provider)
        if not monitor:
            return
        if provider == "docker":
            res = monitor.monitor_containers()
        elif provider == "aws":
            res = monitor.monitor_aws_events()
        else:
            res = monitor.monitor_audit_log()
        self._select_tab_by_title("🎯 Threat Intel")
        self._append_text_widget(
            self.threat_text,
            f"\n[{datetime.now().strftime('%H:%M:%S')}] CLOUD {provider.upper()}: {res.get('status')} - {res.get('reason', '')}\n",
        )

    def _run_framework_report(self, framework):
        data = self.compliance_reporter.generate_framework_report(framework)
        trail = self.compliance_reporter.generate_audit_trail()
        lines = [
            f"📚 FRAMEWORK REPORT ({data['framework'].upper()})",
            f"Generated: {data['generated_at']}",
            f"GRC Score: {data['grc_score']} | Risk: {data['grc_risk']}",
            f"Hygiene Score: {data['hygiene_score']}",
            "",
            "Control Set:",
        ]
        for ctrl in data["controls"]:
            lines.append(f"  - {ctrl}")
        lines.extend([
            "",
            "Audit Trail Summary:",
            f"  Alerts (24h): {trail['alerts_24h']}",
            f"  Config assessments tracked: {trail['config_assessments']}",
            f"  Last GRC run: {trail['grc_assessment']}",
            f"  Last security audit: {trail['security_audit']}",
        ])
        self._set_text_widget(self.results_text, "\n".join(lines))
        self._select_tab_by_title("📊 Scan Results")

    def _handle_dashboard_card_action(self, action):
        actions = {
            "config_assessment": self.config_assessment_gui,
            "malware_detection": self.scan_processes_gui,
            "fim": self.start_fim_monitor_gui,
            "threat_hunting": self.threat_hunt_gui,
            "vuln_detection": self.service_discovery_all_gui,
            "mitre": lambda: self._select_tab_by_title("🎯 Threat Intel"),
            "it_hygiene": self.hygiene_scan_gui,
            "pci": lambda: self._run_framework_report("pci"),
            "gdpr": lambda: self._run_framework_report("gdpr"),
            "hipaa": lambda: self._run_framework_report("hipaa"),
            "docker": lambda: self._show_cloud_status("docker"),
            "aws": lambda: self._show_cloud_status("aws"),
            "gcp": lambda: self._show_cloud_status("aws"),
            "github": lambda: self._show_cloud_status("github"),
        }
        handler = actions.get(action)
        if handler:
            handler()

    def _create_export_tab(self, parent):
        container = tk.Frame(parent, bg='#161b22', padx=30, pady=30)
        container.pack(fill='both', expand=True)

        tk.Label(container, text="📊 Enterprise Reporting & Data Export", 
                 bg='#161b22', fg='#58a6ff', font=("Segoe UI", 16, "bold")).pack(anchor='w', pady=(0, 20))

        # Grid for cards
        grid_frame = tk.Frame(container, bg='#161b22')
        grid_frame.pack(fill='both', expand=True)

        cards = [
            ("📤 Full Text Report", "Generate a comprehensive summary of all network findings in plaintext.", self.generate_report, "#58a6ff"),
            ("💾 Save TXT File", "Save the current summary report to a local text file.", self.save_report_txt_gui, "#58a6ff"),
            ("📦 Export All Data", "Bundle logs, inventory, traffic, and details into a single directory.", self.save_all_data_gui, "#3fb950"),
            ("📄 Export PDF", "Generate a professional PDF document with device details and activity.", self.generate_pdf_report, "#ff7b72"),
            ("📝 Export Word (DOCX)", "Create a Word document including inventory tables and activity logs.", self.generate_docx_report, "#4299e1"),
            ("📊 Devices (CSV)", "Export the current device inventory list to a CSV spreadsheet.", self.export_csv, "#d29922"),
            ("📚 Compliance Report", "Run a framework-specific (PCI/HIPAA/GDPR) compliance audit.", self.generate_framework_report_gui, "#8b949e"),
            ("⚙️ UI Settings", "Adjust window dimensions and layout settings.", self.ui_settings_gui, "#8b949e"),
        ]

        for i, (title, desc, cmd, color) in enumerate(cards):
            row, col = divmod(i, 2)
            card = tk.Frame(grid_frame, bg='#21262d', padx=15, pady=15, highlightthickness=1, highlightbackground="#30363d")
            card.grid(row=row, column=col, sticky='nsew', padx=10, pady=10)
            grid_frame.columnconfigure(col, weight=1)

            lbl_title = tk.Label(card, text=title, bg='#21262d', fg=color, font=("Segoe UI", 12, "bold"))
            lbl_title.pack(anchor='w')
            
            lbl_desc = tk.Label(card, text=desc, bg='#21262d', fg='#8b949e', font=("Segoe UI", 10), 
                               wraplength=400, justify='left')
            lbl_desc.pack(anchor='w', pady=5)
            
            btn = ttk.Button(card, text="Run Export", command=cmd)
            btn.pack(anchor='e', pady=(10, 0))
            
            # Make the card clickable
            for w in (card, lbl_title, lbl_desc):
                w.bind("<Button-1>", lambda _e, c=cmd: c(), add="+")
                w.config(cursor="hand2")

    def alert_timeline_view(self):
        buckets = defaultdict(int)
        for alert in self.alert_manager.alerts:
            stamp = alert["timestamp"].strftime("%Y-%m-%d %H:00")
            buckets[stamp] += 1
        if not buckets:
            return "No alert timeline data yet."
        lines = ["ALERT TIMELINE (hourly):"]
        for hour in sorted(buckets.keys())[-24:]:
            lines.append(f"  {hour}: {buckets[hour]}")
        return "\n".join(lines)

    def alert_correlation_engine(self):
        grouped = defaultdict(list)
        for alert in self.alert_manager.alerts[-500:]:
            grouped[(alert["source_ip"], alert["severity"])].append(alert)
        related = []
        for (ip, sev), alerts in grouped.items():
            if len(alerts) >= 3:
                related.append(f"{ip} ({sev}) has {len(alerts)} correlated alerts")
        return related or ["No strong correlations detected."]

    def trend_analysis(self):
        counts = self.alert_manager.counts_last_24h()
        total = sum(counts.values())
        if total == 0:
            return "No recent alerts for trend analysis."
        dominant = max(counts, key=counts.get)
        return f"Dominant alert trend: {dominant} ({counts[dominant]}/{total} in last 24h)"

    def sli_slo_tracking(self):
        counts = self.alert_manager.counts_last_24h()
        critical_budget = max(0, 5 - counts.get("critical", 0))
        high_budget = max(0, 20 - counts.get("high", 0))
        return (
            "SLI/SLO (24h):\n"
            f"  Critical alert budget remaining: {critical_budget}/5\n"
            f"  High alert budget remaining: {high_budget}/20"
        )
    
    def _create_gui(self):
        self._setup_master_scroll()

        # Update geometry
        self.root.geometry(f"{self.ui_width.get()}x{self.ui_height.get()}")
        
        # Header
        header_frame = tk.Frame(self.app_viewport, bg='#0d1117')
        header_frame.pack(fill='x', pady=(0, 5))

        tk.Label(
            header_frame,
            text="🛡️ NETMON PRO v2.1 - Enterprise Security Operations",
            bg="#0d1117",
            fg="#58a6ff",
            font=("Segoe UI", 16, "bold"),
        ).pack(pady=(5, 2))
        tk.Label(
            header_frame,
            text="Asset Monitoring | Threat Detection | Compliance Auditing",
            bg="#0d1117",
            fg="#8b949e",
            font=("Segoe UI", 10),
        ).pack()
        
        # Main container using grid for precise panel management
        self.main_frame = tk.Frame(self.app_viewport, bg='#161b22')
        self.main_frame.pack(fill='both', expand=True, padx=10, pady=5)
        self.main_frame.columnconfigure(0, weight=0, minsize=self.left_panel_width.get())
        self.main_frame.columnconfigure(1, weight=1)
        self.main_frame.columnconfigure(2, weight=0, minsize=self.right_panel_width.get())
        self.main_frame.rowconfigure(0, weight=1)
        
        # Left panel (controls)
        self.left_panel = tk.Frame(self.main_frame, bg='#161b22')
        self.left_panel.grid(row=0, column=0, sticky='nsew', padx=(0, 10))
        
        # Center panel
        self.center_panel = tk.Frame(self.main_frame, bg='#161b22')
        self.center_panel.grid(row=0, column=1, sticky='nsew', padx=(0, 10))
        
        # Right panel
        self.right_panel = tk.Frame(self.main_frame, bg='#161b22')
        self.right_panel.grid(row=0, column=2, sticky='nsew')
        
        self._create_left_controls(self.left_panel)
        self._create_center_dashboard(self.center_panel)
        self._create_right_details(self.right_panel)

        self.progress_bar = ttk.Progressbar(self.app_viewport, mode="indeterminate")
        self.progress_bar.pack(fill="x", padx=20, pady=(0, 4))

        self.footer_label = tk.Label(
            self.app_viewport,
            text="Ready",
            anchor="w",
            bg="#0d1117",
            fg="#8b949e",
            padx=12,
        )
        self.footer_label.pack(fill="x")
        self._bind_master_wheel_targets(self.app_viewport)
    
    def _create_left_controls(self, parent):
        net_frame = tk.LabelFrame(parent, text="🌐 Network Operations", 
                                 bg='#161b22', fg='#58a6ff', padx=15, pady=15)
        net_frame.pack(fill='x', pady=10)
        
        ttk.Button(net_frame, text="🔍 Discover Router", command=self.discover_router_gui).pack(fill='x', pady=3)
        ttk.Button(net_frame, text="🏠 Scan Local /24", command=self.scan_local_gui).pack(fill='x', pady=3)
        ttk.Button(net_frame, text="👥 Live ARP Monitor", command=self.arp_monitor_gui).pack(fill='x', pady=3)
        
        if NMAP_AVAILABLE:
            ttk.Button(net_frame, text="🌐 Nmap Full Scan", command=self.nmap_scan_gui).pack(fill='x', pady=3)
        
        ttk.Button(net_frame, text="🛠️ Service Discovery (All)", command=self.service_discovery_all_gui).pack(fill='x', pady=3)
        ttk.Button(net_frame, text="🧭 Configuration Assessment", command=self.config_assessment_gui).pack(fill='x', pady=3)
        ttk.Button(net_frame, text="🧬 Baseline FIM", command=self.baseline_filesystem_gui).pack(fill='x', pady=3)
        ttk.Button(net_frame, text="👁️ Start FIM Monitor", command=self.start_fim_monitor_gui).pack(fill='x', pady=3)
        ttk.Button(net_frame, text="🛑 Stop FIM Monitor", command=self.stop_fim_monitor).pack(fill='x', pady=3)

        mon_frame = tk.LabelFrame(parent, text="📊 Live Monitoring", 
                                 bg='#161b22', fg='#58a6ff', padx=15, pady=15)
        mon_frame.pack(fill='x', pady=10)
        
        self.pause_monitor_btn = ttk.Button(mon_frame, text="⏸️ Pause Active Monitoring", command=self.toggle_monitoring_pause)
        self.pause_monitor_btn.pack(fill='x', pady=3)
        ttk.Button(mon_frame, text="📈 Bandwidth Test", command=self.bandwidth_monitor_gui).pack(fill='x', pady=3)
        if SCAPY_AVAILABLE:
            ttk.Button(mon_frame, text="🌐 Packet Capture", command=self.packet_sniffer_gui).pack(fill='x', pady=3)
        
        pentest_frame = tk.LabelFrame(parent, text="🔥 Pentest & Bank Security", 
                                     bg='#161b22', fg='#ff6b6b', padx=15, pady=15)
        pentest_frame.pack(fill='x', pady=10)
        
        ttk.Button(pentest_frame, text="🐛 Suspicious Processes", command=self.scan_processes_gui).pack(fill='x', pady=3)
        ttk.Button(pentest_frame, text="🎯 Deep Device Scan", command=self.deep_device_scan_gui).pack(fill='x', pady=3)
        ttk.Button(pentest_frame, text="🏦 Bank Security Audit", command=self.security_audit_gui).pack(fill='x', pady=3)
        ttk.Button(pentest_frame, text="📘 GRC Compliance Audit", command=self.grc_audit_gui).pack(fill='x', pady=3)
        ttk.Button(pentest_frame, text="🛡️ Critical Asset Watchdog", command=self.asset_watchdog_gui).pack(fill='x', pady=3)
        ttk.Button(pentest_frame, text="🕵️ ARP Poisoning Detection", command=self.arp_poison_detect_gui).pack(fill='x', pady=3)
        ttk.Button(pentest_frame, text="🎯 Threat Hunt", command=self.threat_hunt_gui).pack(fill='x', pady=3)
        ttk.Button(pentest_frame, text="🧹 IT Hygiene Scan", command=self.hygiene_scan_gui).pack(fill='x', pady=3)
        
        rep_frame = tk.LabelFrame(parent, text="📋 Reports & Export", 
                                 bg='#161b22', fg='#58a6ff', padx=15, pady=10)
        rep_frame.pack(fill='x', pady=5)
        
        ttk.Button(rep_frame, text="📤 Full Report (TEXT)", command=self.generate_report).pack(fill='x', pady=2)
        ttk.Button(rep_frame, text="💾 Save Report (TXT)", command=self.save_report_txt_gui).pack(fill='x', pady=2)
        ttk.Button(rep_frame, text="📦 Save All Data & Logs", command=self.save_all_data_gui).pack(fill='x', pady=2)
        ttk.Button(rep_frame, text="📄 Export PDF Report", command=self.generate_pdf_report).pack(fill='x', pady=2)
        ttk.Button(rep_frame, text="📝 Export DOCX Report", command=self.generate_docx_report).pack(fill='x', pady=2)
        ttk.Button(rep_frame, text="💾 Export Devices CSV", command=self.export_csv).pack(fill='x', pady=2)
        ttk.Button(rep_frame, text="📚 Framework Report", command=self.generate_framework_report_gui).pack(fill='x', pady=2)
        ttk.Button(rep_frame, text="⚙️ UI Settings", command=self.ui_settings_gui).pack(fill='x', pady=2)
        ttk.Button(rep_frame, text="🧹 Clear All Data", command=self.clear_all).pack(fill='x', pady=2)
    
    def _create_center_dashboard(self, parent):
        status_frame = tk.LabelFrame(parent, text=" 📈 LIVE SYSTEM OPERATIONAL DASHBOARD ",
                                    bg='#161b22', fg='#58a6ff', padx=15, pady=8)
        status_frame.pack(fill='x', pady=(0, 5))

        status_container = tk.Frame(status_frame, bg='#161b22')
        status_container.pack(fill='x', pady=5)
        for i in range(5): status_container.columnconfigure(i, weight=1)

        self.status_labels = {}
        metric_font = ('Segoe UI', 11, 'bold')
        metrics = [('CPU', '%'), ('RAM', '%'), ('Disk', '%')]
        for i, (name, unit) in enumerate(metrics):
            lbl = tk.Label(status_container, text=f"{name}: --{unit}",
                          bg='#161b22', fg='#58a6ff', font=metric_font)
            lbl.grid(row=0, column=i, padx=5, sticky='w')
            self.status_labels[name.lower()] = lbl

        self.net_label = tk.Label(status_container, text="Network: --",
                                 bg='#161b22', fg='#58a6ff', font=metric_font)
        self.net_label.grid(row=0, column=3, padx=5, sticky='w')

        self.router_label = tk.Label(status_container, text="Router: --",
                                    bg='#161b22', fg='#58a6ff', font=metric_font)
        self.router_label.grid(row=0, column=4, padx=5, sticky='w')

        self.live_label = tk.Label(status_container, text="LIVE: 0",
                                   bg="#161b22", fg="#3fb950", font=metric_font)
        self.live_label.grid(row=1, column=0, padx=5, pady=5, sticky='w')

        self.critical_label = tk.Label(status_container, text="CRITICAL: 0",
                                       bg="#161b22", fg="#f85149", font=metric_font)
        self.critical_label.grid(row=1, column=1, padx=5, pady=5, sticky='w')

        self.monitor_state_label = tk.Label(status_container, text="MONITOR: RUNNING",
                                            bg="#161b22", fg="#3fb950", font=metric_font)
        self.monitor_state_label.grid(row=1, column=3, columnspan=2, padx=5, pady=5, sticky='e')

        self.devices_label = tk.Label(status_container, text="Devices: 0",
                                     bg='#161b22', fg='#58a6ff', font=metric_font)
        self.devices_label.grid(row=1, column=2, padx=5, pady=5, sticky='w')

        grc_kpi = tk.Frame(status_frame, bg="#0d1117", relief="solid", bd=1)
        grc_kpi.pack(fill="x", pady=(5, 0))
        self.grc_score_label = tk.Label(grc_kpi, text="GRC Score: --", bg="#0d1117", fg="#58a6ff", font=("Segoe UI", 10, "bold"))
        self.grc_score_label.pack(side="left", padx=10, pady=5)
        self.grc_risk_label = tk.Label(grc_kpi, text="Risk Level: --", bg="#0d1117", fg="#ff6b6b", font=("Segoe UI", 10, "bold"))
        self.grc_risk_label.pack(side="left", padx=10, pady=5)
        self.grc_framework_label = tk.Label(grc_kpi, text="NIST | ISO | PCI | CIS", bg="#0d1117", fg="#8b949e", font=("Segoe UI", 10, "bold"))
        self.grc_framework_label.pack(side="right", padx=10, pady=5)

        self.notebook = ttk.Notebook(parent)
        self.notebook.pack(fill='both', expand=True, padx=10, pady=8)

        log_frame = tk.Frame(self.notebook, bg='#0d1117')
        self.notebook.add(log_frame, text="Operations\nLog")
        self.log_text = scrolledtext.ScrolledText(log_frame, bg='#0d1117', fg='#c9d1d9',
                                                 font=('Consolas', 11), wrap=tk.WORD, height=20)
        self.log_text.pack(fill='both', expand=True, padx=10, pady=10)

        self.log_text.tag_config('INFO', foreground='#58a6ff')
        self.log_text.tag_config('WARN', foreground='#d29922')
        self.log_text.tag_config('ERROR', foreground='#f85149')
        self.log_text.tag_config('DEBUG', foreground='#8b949e')
        self.log_text.tag_config('CRITICAL', foreground='#ff6b6b', font=('Consolas', 11, 'bold'))
        self.log_text.tag_config('SUCCESS', foreground='#3fb950')

        devices_frame = tk.Frame(self.notebook, bg='#0d1117')
        self.notebook.add(devices_frame, text="Device\nInventory")
        filter_frame = tk.Frame(devices_frame, bg="#0d1117")
        filter_frame.pack(fill="x", padx=10, pady=(10, 0))
        tk.Label(filter_frame, text="Filter:", bg="#0d1117", fg="#8b949e", font=("Segoe UI", 10, "bold")).pack(side="left")
        ttk.Entry(filter_frame, textvariable=self.device_filter_var).pack(side="left", fill="x", expand=True, padx=(8, 6))
        ttk.Checkbutton(filter_frame, text="LIVE only", variable=self.live_only_var, command=self._update_devices_tree).pack(side="left", padx=(0, 6))
        ttk.Button(filter_frame, text="Clear", command=lambda: self.device_filter_var.set("")).pack(side="left")
        self.device_filter_var.trace_add("write", lambda *_: self._update_devices_tree())

        tree_frame = tk.Frame(devices_frame)
        tree_frame.pack(fill='both', expand=True, padx=10, pady=(8, 5))
        self.devices_tree = ttk.Treeview(tree_frame, columns=('IP', 'MAC', 'NAME', 'HOSTNAME', 'VENDOR', 'STATUS'),
                                         show='headings', height=18)
        self._setup_devices_tree()
        tree_scroll = ttk.Scrollbar(tree_frame, orient='vertical', command=self.devices_tree.yview)
        self.devices_tree.configure(yscrollcommand=tree_scroll.set)
        self.devices_tree.pack(side='left', fill='both', expand=True)
        tree_scroll.pack(side='right', fill='y')
        self.devices_tree.bind('<<TreeviewSelect>>', self._on_device_select)

        tab_specs = [
            ("Scan\nResults", "results_text"),
            ("Alert\nCenter", "alerts_text"),
            ("Activity\nMonitor", "activity_text"),
            ("Endpoint\nSecurity", "endpoint_text"),
            ("File\nIntegrity", "fim_text"),
            ("Threat\nIntel", "threat_text"),
            ("Network\nServices", "services_text"),
            ("Security\nAudit", "audit_text"),
            ("GRC\nCompliance", "grc_text"),
        ]

        self.content_text_widgets = [self.log_text]
        for title, attr in tab_specs:
            frame = tk.Frame(self.notebook, bg='#0d1117')
            self.notebook.add(frame, text=title)
            widget = scrolledtext.ScrolledText(frame, bg='#0d1117', fg='#c9d1d9',
                                               font=('Consolas', 11), wrap=tk.WORD)
            widget.pack(fill='both', expand=True, padx=10, pady=10)
            setattr(self, attr, widget)
            self.content_text_widgets.append(widget)

        export_tab = tk.Frame(self.notebook, bg='#161b22')
        self.notebook.add(export_tab, text="Export\nReports")
        self._create_export_tab(export_tab)

        for w in self.content_text_widgets:
            self._enable_smooth_scroll(w)
        self._set_content_display_mode()

    def _create_alert_stat(self, parent, title, value, color, subtitle, on_click=None):
        box = tk.Frame(parent, bg="#ffffff")
        box.pack(side="left", expand=True, fill="x", padx=6)
        title_lbl = tk.Label(box, text=title, bg="#ffffff", fg="#2d3748", font=("Segoe UI", 11))
        title_lbl.pack()
        value_lbl = tk.Label(box, text=value, bg="#ffffff", fg=color, font=("Segoe UI", 30))
        value_lbl.pack()
        sub_lbl = tk.Label(box, text=subtitle, bg="#ffffff", fg="#6b7280", font=("Segoe UI", 10))
        sub_lbl.pack()
        if on_click:
            for w in (box, title_lbl, value_lbl, sub_lbl):
                w.bind("<Button-1>", lambda _e: on_click(), add="+")
                w.config(cursor="hand2")
        return value_lbl

    def _create_section_panel(self, parent, row, col, title, cards):
        section = tk.Frame(parent, bg="#ffffff", bd=1, relief="solid")
        section.grid(row=row, column=col, sticky="nsew", padx=5, pady=6)
        tk.Label(section, text=title, bg="#ffffff", fg="#4a5568", font=("Segoe UI", 10, "bold")).pack(pady=(8, 6))

        cards_wrap = tk.Frame(section, bg="#ffffff")
        cards_wrap.pack(fill="both", expand=True, padx=8, pady=(0, 8))
        columns = 2
        clickable_cards = [c for c in cards if c.get("action")]
        for idx, card_meta in enumerate(clickable_cards):
            card_title = card_meta["title"]
            card_desc = card_meta["desc"]
            card_action = card_meta["action"]
            r = idx // columns
            c = idx % columns
            card = tk.Frame(cards_wrap, bg="#f8fafc", bd=1, relief="solid", cursor="hand2")
            card.grid(row=r, column=c, sticky="nsew", padx=5, pady=5)
            cards_wrap.grid_columnconfigure(c, weight=1)
            title_lbl = tk.Label(card, text=card_title, bg="#f8fafc", fg="#2d3748", font=("Segoe UI", 11, "bold"))
            title_lbl.pack(anchor="w", padx=10, pady=(10, 4))
            desc_lbl = tk.Label(card, text=card_desc, bg="#f8fafc", fg="#4a5568", font=("Segoe UI", 10), justify="left", wraplength=320)
            desc_lbl.pack(anchor="w", padx=10, pady=(0, 10))
            for w in (card, title_lbl, desc_lbl):
                w.bind("<Button-1>", lambda _e, a=card_action: self._handle_dashboard_card_action(a), add="+")

    def _draw_agents_donut(self, active_count, disconnected_count):
        if not hasattr(self, "agents_canvas"):
            return
        canvas = self.agents_canvas
        canvas.delete("all")
        x1, y1, x2, y2 = 15, 15, 135, 135
        total = max(active_count + disconnected_count, 1)
        active_extent = (active_count / total) * 360
        canvas.create_oval(x1, y1, x2, y2, outline="#e2e8f0", width=14)
        if active_count > 0:
            canvas.create_arc(x1, y1, x2, y2, start=90, extent=-active_extent, style="arc", outline="#0f766e", width=14)
        if disconnected_count > 0:
            canvas.create_arc(x1, y1, x2, y2, start=90 - active_extent, extent=-(360 - active_extent), style="arc", outline="#c53030", width=14)

    def _update_overview_dashboard(self):
        with self.state_lock:
            total_devices = len(self.devices)
            active_count = sum(1 for info in self.devices.values() if "LIVE" in info.get("status", ""))
            disconnected_count = max(total_devices - active_count, 0)
        counts = self.alert_manager.counts_last_24h()
        critical_count = counts.get("critical", 0)
        high_count = counts.get("high", 0)
        medium_count = counts.get("medium", 0)
        low_count = counts.get("low", 0)

        if hasattr(self, "active_agents_label"):
            self.active_agents_label.config(text=f"● Active ({active_count})")
        if hasattr(self, "disconnected_agents_label"):
            self.disconnected_agents_label.config(text=f"● Disconnected ({disconnected_count})")
        if hasattr(self, "critical_severity_label"):
            self.critical_severity_label.config(text=str(critical_count))
        if hasattr(self, "high_severity_label"):
            self.high_severity_label.config(text=str(high_count))
        if hasattr(self, "medium_severity_label"):
            self.medium_severity_label.config(text=str(medium_count))
        if hasattr(self, "low_severity_label"):
            self.low_severity_label.config(text=str(low_count))
        self._draw_agents_donut(active_count, disconnected_count)
    
    def _setup_devices_tree(self):
        self.devices_tree.heading('IP', text='IP Address')
        self.devices_tree.heading('MAC', text='MAC Address')
        self.devices_tree.heading('NAME', text='Custom Name')
        self.devices_tree.heading('HOSTNAME', text='Hostname')
        self.devices_tree.heading('VENDOR', text='Vendor')
        self.devices_tree.heading('STATUS', text='Status')
        
        self.devices_tree.column('IP', width=130, anchor='center')
        self.devices_tree.column('MAC', width=150, anchor='center')
        self.devices_tree.column('NAME', width=140, anchor='w')
        self.devices_tree.column('HOSTNAME', width=140, anchor='w')
        self.devices_tree.column('VENDOR', width=160, anchor='w')
        self.devices_tree.column('STATUS', width=100, anchor='center')
    
    def _create_right_details(self, parent):
        details_frame = tk.LabelFrame(parent, text="🎯 ASSET INTELLIGENCE", 
                                     bg='#161b22', fg='#58a6ff', padx=10, pady=10)
        details_frame.pack(fill='both', expand=True, pady=(0, 5))
        
        summary_frame = tk.Frame(details_frame, bg='#0d1117', relief='raised', bd=1)
        summary_frame.pack(fill='x', pady=(0, 5))
        
        self.summary_label = tk.Label(summary_frame, text="No Selection", font=('Segoe UI', 12, 'bold'),
                                     bg='#0d1117', fg='#58a6ff')
        self.summary_label.pack(pady=10)
        
        self.summary_text = tk.Text(summary_frame, height=3, bg='#0d1117', fg='#c9d1d9',
                                   font=('Consolas', 10), wrap=tk.WORD, state='disabled')
        self.summary_text.pack(fill='x', padx=10, pady=(0, 10))
        self._enable_smooth_scroll(self.summary_text)
        
        name_frame = tk.Frame(details_frame, bg='#161b22')
        name_frame.pack(fill='x', pady=2)
        tk.Label(name_frame, text="Alias:", bg='#161b22', fg='#c9d1d9', font=('Segoe UI', 9)).pack(side='left')
        self.custom_name_entry = ttk.Entry(name_frame)
        self.custom_name_entry.pack(side='left', fill='x', expand=True, padx=5)
        ttk.Button(name_frame, text="Set", width=5, command=self.set_custom_name).pack(side='left')

        # High-density buttons
        btn_grid = tk.Frame(details_frame, bg='#161b22')
        btn_grid.pack(fill='x', pady=5)
        
        ops = [
            ("🔍 Deep Scan", self.deep_device_scan_gui),
            ("🌐 Nmap", self.nmap_scan_selected_gui),
            ("📡 Route", self.traceroute_selected),
            ("🛡️ Critical", self.toggle_critical_asset),
            ("💾 Save", self.save_device_details),
            ("🗑️ Clear", self.clear_saved_device_gui)
        ]
        
        for i, (text, cmd) in enumerate(ops):
            row, col = divmod(i, 2)
            btn = ttk.Button(btn_grid, text=text, command=cmd)
            btn.grid(row=row, column=col, sticky='ew', padx=1, pady=1)
        btn_grid.columnconfigure(0, weight=1)
        btn_grid.columnconfigure(1, weight=1)
        
        self.details_text = scrolledtext.ScrolledText(details_frame, bg='#0d1117', fg='#c9d1d9',
                                                     font=('Consolas', 10), wrap=tk.WORD, height=15)
        self.details_text.pack(fill='both', expand=True, padx=2, pady=5)
        self._enable_smooth_scroll(self.details_text)
    
    def _on_device_select(self, event):
        selection = self.devices_tree.selection()
        if selection:
            item = self.devices_tree.item(selection[0])
            ip = item['values'][0]
            self.selected_device = ip
            self._show_device_details(ip)
    
    def set_custom_name(self):
        if not self.selected_device: return
        name = self._sanitize_label(self.custom_name_entry.get())
        if self.selected_device in self.devices:
            self.devices[self.selected_device]['custom_name'] = name
            self._save_to_db(self.selected_device, self.devices[self.selected_device])
            self._update_devices_tree()
            self.log(f"🏷️ Set custom name for {self.selected_device} to '{name}'")

    def _set_monitor_pause_ui(self):
        if self.monitoring_paused:
            if hasattr(self, "pause_monitor_btn"):
                self.pause_monitor_btn.config(text="▶️ Resume Active Monitoring")
            if hasattr(self, "monitor_state_label"):
                self.monitor_state_label.config(text="MONITOR: PAUSED", fg="#d29922")
        else:
            if hasattr(self, "pause_monitor_btn"):
                self.pause_monitor_btn.config(text="⏸️ Pause Active Monitoring")
            if hasattr(self, "monitor_state_label"):
                self.monitor_state_label.config(text="MONITOR: RUNNING", fg="#3fb950")

    def toggle_monitoring_pause(self):
        self.monitoring_paused = not self.monitoring_paused
        self._set_monitor_pause_ui()
        if self.monitoring_paused:
            self._set_footer("Active monitoring paused")
            self.log("⏸️ Active monitoring paused")
        else:
            self._set_footer("Active monitoring resumed")
            self.log("▶️ Active monitoring resumed")

    def clear_saved_device_gui(self):
        if not self.selected_device:
            messagebox.showwarning("No Selection", "Select an IP/device from inventory first.")
            return
        ip = self.selected_device
        if not messagebox.askyesno("Clear Saved Device", f"Delete saved device data for {ip}?"):
            return
        self._clear_saved_device(ip)

    def _clear_saved_device(self, ip):
        with self.state_lock:
            removed = self.devices.pop(ip, None)
            self.device_details.pop(ip, None)
            self.traffic_logs.pop(ip, None)
        if removed is None:
            messagebox.showinfo("Not Found", f"No saved entry found for {ip}.")
            return
        try:
            with sqlite3.connect(DB_FILE) as conn:
                conn.execute("DELETE FROM devices WHERE ip = ?", (ip,))
        except Exception as exc:
            self.log(f"⚠️ Failed to remove {ip} from history DB: {exc}", "WARN")
        if self.selected_device == ip:
            self.selected_device = None
            self.summary_label.config(text="📱 No device selected")
            self.custom_name_entry.delete(0, tk.END)
            self.summary_text.config(state='normal')
            self.summary_text.delete(1.0, tk.END)
            self.summary_text.config(state='disabled')
            self._set_text_widget(self.details_text, "")
        self._update_devices_tree()
        self._update_overview_dashboard()
        self.log(f"🗑️ Cleared saved device: {ip}")

    def ui_settings_gui(self):
        settings_win = tk.Toplevel(self.root)
        settings_win.title("⚙️ UI Settings")
        settings_win.geometry("560x580")
        settings_win.configure(bg='#161b22')

        tk.Label(settings_win, text="Window Width:", bg='#161b22', fg='#c9d1d9').pack(pady=5)
        tk.Scale(settings_win, from_=1000, to=2500, orient='horizontal', variable=self.ui_width, bg='#161b22', fg='#c9d1d9').pack(fill='x', padx=20)
        
        tk.Label(settings_win, text="Window Height:", bg='#161b22', fg='#c9d1d9').pack(pady=5)
        tk.Scale(settings_win, from_=600, to=1500, orient='horizontal', variable=self.ui_height, bg='#161b22', fg='#c9d1d9').pack(fill='x', padx=20)

        tk.Label(settings_win, text="Left Panel Width:", bg='#161b22', fg='#c9d1d9').pack(pady=5)
        tk.Scale(settings_win, from_=200, to=600, orient='horizontal', variable=self.left_panel_width, bg='#161b22', fg='#c9d1d9').pack(fill='x', padx=20)

        tk.Label(settings_win, text="Right Panel Width:", bg='#161b22', fg='#c9d1d9').pack(pady=5)
        tk.Scale(settings_win, from_=200, to=700, orient='horizontal', variable=self.right_panel_width, bg='#161b22', fg='#c9d1d9').pack(fill='x', padx=20)

        tk.Label(settings_win, text="FIM Polling Interval (seconds):", bg="#161b22", fg="#c9d1d9").pack(pady=(12, 4))
        tk.Scale(settings_win, from_=5, to=120, orient="horizontal", variable=self.fim_interval, bg="#161b22", fg="#c9d1d9").pack(fill="x", padx=20)

        fim_path_row = tk.Frame(settings_win, bg="#161b22")
        fim_path_row.pack(fill="x", padx=20, pady=(8, 6))
        tk.Label(fim_path_row, text="Default FIM Path:", bg="#161b22", fg="#c9d1d9").pack(side="left")
        ttk.Entry(fim_path_row, textvariable=self.default_fim_path).pack(side="left", fill="x", expand=True, padx=8)

        ttk.Checkbutton(settings_win, text="Enable Auto-Remediation (safe actions)", variable=self.auto_remediate).pack(anchor="w", padx=20, pady=(6, 0))

        ttk.Button(settings_win, text="Auto-Fit To Screen", command=self._autofit_ui).pack(pady=(14, 4))
        ttk.Button(settings_win, text="Apply Changes", command=self._apply_ui_changes).pack(pady=20)

    def _autofit_ui(self):
        screen_w = self.root.winfo_screenwidth()
        screen_h = self.root.winfo_screenheight()
        fit_w = min(1800, max(1400, screen_w - 40))
        fit_h = min(1100, max(900, screen_h - 80))
        self.ui_width.set(fit_w)
        self.ui_height.set(fit_h)
        self.left_panel_width.set(min(400, max(260, int(fit_w * 0.22))))
        self.right_panel_width.set(min(450, max(280, int(fit_w * 0.24))))
        self._apply_ui_changes()
        self.log("🖥️ UI auto-fit applied for current screen")

    def _apply_ui_changes(self):
        self.root.geometry(f"{self.ui_width.get()}x{self.ui_height.get()}")
        self.left_panel.config(width=self.left_panel_width.get())
        self.right_panel.config(width=self.right_panel_width.get())
        self._set_content_display_mode()
        self.log("🎨 UI updated")

    def _set_content_display_mode(self):
        text_font = ("Consolas", 11)
        text_padx = 10
        text_pady = 10

        for widget in getattr(self, "content_text_widgets", ()):
            try:
                widget.config(font=text_font, wrap=tk.WORD)
                widget.pack_configure(padx=text_padx, pady=text_pady)
            except Exception:
                continue

        if hasattr(self, "notebook"):
            self.notebook.pack_configure(padx=10, pady=8)
        if hasattr(self, "devices_tree"):
            style = ttk.Style()
            style.configure("Treeview", rowheight=26)
            self._setup_devices_tree()

    def _show_device_details(self, ip):
        with self.state_lock:
            if ip not in self.devices:
                return
            device = dict(self.devices[ip])
            details = dict(self.device_details.get(ip, {}))
            traffic = list(self.traffic_logs.get(ip, []))
        
        self.summary_label.config(text=f"📱 {ip}")
        custom = device.get('custom_name', '')
        self.custom_name_entry.delete(0, tk.END)
        self.custom_name_entry.insert(0, custom)

        summary = f"MAC: {device.get('mac', 'Unknown')}\nHostname: {device.get('hostname', 'Unknown')}\nName: {custom or 'N/A'}\nStatus: {device.get('status', 'Unknown')}"
        
        self.summary_text.config(state='normal')
        self.summary_text.delete(1.0, tk.END)
        self.summary_text.insert(tk.END, summary)
        self.summary_text.config(state='disabled')
        
        full_details = f"""🎯 DEVICE INTEL: {ip} ({custom or 'No Name'})
═══════════════════════════════════════
🔗 BASIC INFO
IP:           {ip}
MAC:          {device.get('mac', 'N/A')}
Hostname:     {device.get('hostname', 'N/A')}
Vendor:       {device.get('vendor', 'N/A')}
OS Detected:  {details.get('os', 'Unknown')}
Connected:    {device.get('connection_time', 'N/A')}

📊 DISCOVERY DATA
First Seen:   {details.get('first_seen', 'N/A')}
Last Seen:    {details.get('last_seen', 'N/A')}
Open Ports:   {len(details.get('ports', {}))}

🔍 SERVICES
{chr(10).join([f"  • {p} ({s})" for p, s in details.get('ports', {}).items()]) or 'None'}

🌐 TRAFFIC LOGS (DNS)
{chr(10).join(traffic[-10:]) or 'No traffic recorded'}

🌐 NETWORK PATH
{details.get('traceroute', 'Not scanned')}
"""
        self.details_text.delete(1.0, tk.END)
        self.details_text.insert(tk.END, full_details)
    
    def _start_monitoring(self):
        def update_status():
            if self.monitoring_paused:
                self.root.after(2000, update_status)
                return
            try:
                cpu = psutil.cpu_percent(interval=0.1)
                ram = psutil.virtual_memory().percent
                disk = psutil.disk_usage('/').percent
                net = psutil.net_io_counters()
                
                self.status_labels['cpu'].config(text=f"CPU: {cpu:.1f}%")
                self.status_labels['ram'].config(text=f"RAM: {ram:.1f}%")
                self.status_labels['disk'].config(text=f"Disk: {disk:.1f}%")
                
                up = net.bytes_sent / 1024 / 1024
                down = net.bytes_recv / 1024 / 1024
                self.net_label.config(text=f"↑{up:.1f}MB ↓{down:.1f}MB")
                with self.state_lock:
                    all_devices = len(self.devices)
                    live_devices = sum(1 for info in self.devices.values() if "LIVE" in info.get("status", ""))
                    critical_devices = sum(1 for info in self.devices.values() if info.get("is_critical", False))
                self.devices_label.config(text=f"Devices: {all_devices}")
                self.live_label.config(text=f"LIVE: {live_devices}")
                self.critical_label.config(text=f"Critical: {critical_devices}")
                self._update_overview_dashboard()
            except: pass
            self.root.after(2000, update_status)
        update_status()

    def _auto_discover_network(self):
        threading.Thread(target=self._discover_network_info, daemon=True).start()
    
    def _discover_network_info(self):
        try:
            s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
            s.connect(("8.8.8.8", 80))
            local_ip = s.getsockname()[0]
            s.close()
            
            try:
                result = self._run_command(['ip', 'route', 'get', '1.1.1.1'], capture_output=True, timeout=5, silent=True)
                if result is None:
                    raise RuntimeError("ip command unavailable")
                for line in result.stdout.split('\n'):
                    if 'via' in line:
                        self.router_ip = line.split('via')[1].split()[0]
                        break
            except: self.router_ip = "192.168.1.1"
            
            network = ipaddress.IPv4Network(f"{local_ip}/24", strict=False)
            self.isp_range = list(network.hosts())
            if self.router_ip:
                self._queue_ui(self.router_label.config, text=f"Router: {self.router_ip}")
            self.log(f"✅ Network: {local_ip}/24 discovered")
        except Exception as e:
            self.log(f"❌ Discovery error: {e}", "ERROR")

    def discover_router_gui(self):
        threading.Thread(target=self._discover_router, daemon=True).start()
    
    def _discover_router(self):
        self.log("🔍 Router discovery...")
        self._queue_ui(self._start_task_ui, "Discovering router...")
        
        # Try to get from system route first
        try:
            if platform.system() == 'Windows':
                res = self._run_command(['route', 'print', '0.0.0.0'], capture_output=True, timeout=5, silent=True)
                if res and res.stdout:
                    match = re.search(r'0\.0\.0\.0\s+0\.0\.0\.0\s+(\d+\.\d+\.\d+\.\d+)', res.stdout)
                    if match:
                        gw = match.group(1)
                        self.router_ip = gw
                        self._queue_ui(self.router_label.config, text=f"Router: {gw} ✅")
                        self.log(f"✅ Router discovered via routing table: {gw}")
                        self._queue_ui(self._end_task_ui)
                        return
            else:
                res = self._run_command(['ip', 'route', 'show', 'default'], capture_output=True, timeout=5, silent=True)
                if res and res.stdout:
                    match = re.search(r'default via (\d+\.\d+\.\d+\.\d+)', res.stdout)
                    if match:
                        gw = match.group(1)
                        self.router_ip = gw
                        self._queue_ui(self.router_label.config, text=f"Router: {gw} ✅")
                        self.log(f"✅ Router discovered via ip route: {gw}")
                        self._queue_ui(self._end_task_ui)
                        return
        except: pass

        gateways = ['192.168.1.1', '192.168.0.1', '10.0.0.1', '172.16.0.1', '172.16.1.1']
        ping_cmd = ['ping', '-n', '1', '-w', '2000'] if platform.system() == 'Windows' else ['ping', '-c', '1', '-W', '2']
        for gw in gateways:
            try:
                result = self._run_command([*ping_cmd, gw], timeout=4, silent=True)
                if result and result.returncode == 0:
                    self.router_ip = gw
                    self._queue_ui(self.router_label.config, text=f"Router: {gw} ✅")
                    self._queue_ui(self._end_task_ui)
                    self.log(f"✅ Router LIVE (fallback scan): {gw}")
                    return
            except: continue
        self._queue_ui(self._end_task_ui)
        self.log("⚠️ No router ping response", "WARN")

    def scan_local_gui(self):
        if not self.isp_range:
            self.log("⚠️ Network range unavailable. Discovering network first...", "WARN")
            self._auto_discover_network()
            return
        threading.Thread(target=self._scan_local, daemon=True).start()
    
    def _scan_local(self):
        self.log("🏠 Scanning local /24 network...")
        self._queue_ui(self._start_task_ui, "Scanning local /24...")
        live_count = 0
        ping_cmd = ['ping', '-n', '1', '-w', '1000'] if platform.system() == 'Windows' else ['ping', '-c', '1', '-W', '1']
        for i in range(min(254, len(self.isp_range))):
            ip = str(self.isp_range[i])
            try:
                result = self._run_command([*ping_cmd, ip], timeout=2, silent=True)
                if result and result.returncode == 0:
                    live_count += 1
                    resolved_host = self._resolve_hostname(ip)
                    discovered_mac = self._get_mac_for_ip(ip)
                    with self.state_lock:
                        info = self.devices.get(ip, {
                            'mac': '??:??:??:??:??:??', 
                            'hostname': 'Unknown',
                            'vendor': 'Unknown', 
                            'status': 'LIVE',
                            'connection_time': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                        })
                        current_mac = self._normalize_mac(info.get('mac', ''))
                        if self._is_unknown_identity(current_mac) and discovered_mac:
                            current_mac = discovered_mac

                        current_host = info.get('hostname', 'Unknown')
                        if self._is_unknown_identity(current_host) and not self._is_unknown_identity(resolved_host):
                            current_host = resolved_host

                        current_vendor = info.get('vendor', 'Unknown')
                        if not self._is_unknown_identity(current_mac):
                            looked_up_vendor = self._get_vendor(current_mac)
                            if not self._is_unknown_identity(looked_up_vendor):
                                current_vendor = looked_up_vendor

                        info['mac'] = current_mac or '??:??:??:??:??:??'
                        info['hostname'] = current_host if current_host else 'Unknown'
                        info['vendor'] = current_vendor if current_vendor else 'Unknown'
                        info['status'] = 'LIVE'
                        if 'connection_time' not in info or not info['connection_time']:
                            info['connection_time'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                        self.devices[ip] = info
                    self.device_details[ip] = {'first_seen': datetime.now().strftime("%H:%M:%S")}
                    self._save_to_db(ip, info)
                    self._queue_ui(self._update_devices_tree)
            except: continue
        self._queue_ui(self._end_task_ui)
        self.log(f"✅ Local scan: {live_count} LIVE hosts")

    def _resolve_hostname(self, ip):
        try: return socket.gethostbyaddr(ip)[0]
        except: return "Unknown"

    def _is_unknown_identity(self, value):
        text = str(value or "").strip().lower()
        return (not text) or text in {"unknown", "n/a", "none", "null"} or "??" in text

    def _normalize_mac(self, mac):
        value = str(mac or "").strip().lower().replace("-", ":")
        if not value:
            return ""
        if re.fullmatch(r"[0-9a-f]{12}", value):
            value = ":".join(value[i:i + 2] for i in range(0, 12, 2))
        if re.fullmatch(r"(?:[0-9a-f]{2}:){5}[0-9a-f]{2}", value):
            return value
        return ""

    def _extract_mac(self, text):
        if not text:
            return ""
        match = re.search(r"((?:[0-9A-Fa-f]{2}[:-]){5}[0-9A-Fa-f]{2})", text)
        return self._normalize_mac(match.group(1)) if match else ""

    def _get_mac_for_ip(self, ip):
        commands = []
        if platform.system() == "Windows":
            commands = [['arp', '-a', ip]]
        else:
            commands = [['ip', 'neigh', 'show', ip], ['arp', '-n', ip]]

        for cmd in commands:
            try:
                res = self._run_command(cmd, capture_output=True, timeout=3, silent=True)
                if not res or not res.stdout:
                    continue
                mac = self._extract_mac(res.stdout)
                if mac:
                    return mac
            except Exception:
                continue
        return ""

    def arp_monitor_gui(self):
        if not SCAPY_AVAILABLE:
            messagebox.showerror("Missing Dependency", "sudo apt install python3-scapy")
            return
        threading.Thread(target=self._arp_monitor, daemon=True).start()
    
    def _arp_monitor(self):
        if sniff is None or ARP is None:
            self.log("❌ ARP monitor unavailable: Scapy symbols not loaded", "ERROR")
            return
        self._queue_ui(self._start_task_ui, "Monitoring ARP packets...")
        def arp_handler(pkt):
            if self.monitoring_paused:
                return
            if ARP in pkt and pkt[ARP].op == 2:
                ip, mac = pkt[ARP].psrc, pkt[ARP].hwsrc
                with self.state_lock:
                    info = self.devices.get(ip, {})
                    info.update({
                        'mac': mac, 
                        'vendor': self._get_vendor(mac), 
                        'status': 'LIVE', 
                        'hostname': self._resolve_hostname(ip)
                    })
                    if 'connection_time' not in info or not info['connection_time']:
                        info['connection_time'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    self.devices[ip] = info
                self._save_to_db(ip, info)
                self._queue_ui(self._update_devices_tree)
        
        self.log("👥 LIVE ARP monitoring (60s)...")
        sniff(prn=arp_handler, filter="arp", timeout=60, store=False)
        self._queue_ui(self._end_task_ui)
        self.log("✅ ARP monitoring complete")

    def _get_vendor(self, mac):
        oui_db = {'00:50:56': 'VMware', '08:00:27': 'VirtualBox', 'B8:27:EB': 'Raspberry Pi'}
        normalized = self._normalize_mac(mac)
        if not normalized:
            return 'Unknown'
        prefix = normalized[:8].upper()
        if prefix in oui_db:
            return oui_db[prefix]

        oui_long = normalized.replace(":", "").upper()[:6]
        if not hasattr(self, "_oui_vendor_map"):
            self._oui_vendor_map = {}
            for path in ("/usr/share/nmap/nmap-mac-prefixes", "/usr/share/ieee-data/oui.txt"):
                if not os.path.exists(path):
                    continue
                try:
                    with open(path, "r", encoding="utf-8", errors="ignore") as f:
                        for line in f:
                            line = line.strip()
                            if not line or line.startswith("#"):
                                continue
                            if path.endswith("nmap-mac-prefixes"):
                                parts = line.split(None, 1)
                                if len(parts) == 2 and re.fullmatch(r"[0-9A-Fa-f]{6}", parts[0]):
                                    self._oui_vendor_map[parts[0].upper()] = parts[1].strip()
                            elif "(base 16)" in line:
                                left, right = line.split("(base 16)", 1)
                                key = re.sub(r"[^0-9A-Fa-f]", "", left).upper()[:6]
                                vendor = right.strip()
                                if len(key) == 6 and vendor:
                                    self._oui_vendor_map[key] = vendor
                except Exception:
                    continue
        return self._oui_vendor_map.get(oui_long, 'Unknown')

    def _update_devices_tree(self):
        search = self.device_filter_var.get().strip().lower()
        live_only = self.live_only_var.get()
        for item in self.devices_tree.get_children():
            self.devices_tree.delete(item)
        visible = 0
        with self.state_lock:
            snapshot = list(self.devices.items())
        for ip, info in sorted(snapshot, key=lambda x: 'LIVE' in x[1].get('status', ''), reverse=True):
            haystack = " ".join([
                ip,
                info.get("mac", ""),
                info.get("custom_name", ""),
                info.get("hostname", ""),
                info.get("vendor", ""),
                info.get("status", ""),
            ]).lower()
            if search and search not in haystack:
                continue
            if live_only and "LIVE" not in info.get("status", ""):
                continue
            status_icon = "🟢" if "LIVE" in info.get('status', '') else "⚪"
            crit_icon = "🛡️ " if info.get('is_critical', False) else ""
            self.devices_tree.insert('', 'end', values=(ip, info.get('mac', 'Unknown'), f"{crit_icon}{info.get('custom_name', '')}", info.get('hostname', ''), info.get('vendor', 'Unknown'), f"{status_icon} {info.get('status')}"))
            visible += 1
            if visible >= 200:
                break

    def deep_device_scan_gui(self):
        if not self.selected_device:
            messagebox.showwarning("No Selection", "Select a device first!")
            return
        if not self._is_valid_ipv4(self.selected_device):
            messagebox.showerror("Invalid IP", f"Cannot scan invalid IP: {self.selected_device}")
            return
        threading.Thread(target=lambda: self._deep_scan_device(self.selected_device), daemon=True).start()

    def config_assessment_gui(self):
        threading.Thread(target=self._run_config_assessment, daemon=True).start()

    def check_compliance_rules(self, ip, info, details):
        findings = []
        os_name = details.get("os", "").lower()
        ports = set(int(p) for p in details.get("ports", {}).keys())
        status = info.get("status", "")

        if "LIVE" in status and not info.get("custom_name", "").strip():
            findings.append({
                "rule": "Asset Naming Standard",
                "severity": "medium",
                "frameworks": ["CIS 1.1", "NIST ID.AM-1"],
                "remediation": "Set a descriptive custom asset name",
            })
        if ports & {21, 23}:
            findings.append({
                "rule": "Insecure Management Protocols (FTP/Telnet)",
                "severity": "high",
                "frameworks": ["PCI DSS Req.2", "CIS 4"],
                "remediation": "Disable FTP/Telnet and enforce SSH/SFTP",
            })
        if 3389 in ports:
            findings.append({
                "rule": "RDP Exposure",
                "severity": "high",
                "frameworks": ["NIST PR.AC", "HIPAA Access Control"],
                "remediation": "Restrict RDP via VPN/MFA and ACLs",
            })
        if 445 in ports:
            findings.append({
                "rule": "SMB Exposure",
                "severity": "high",
                "frameworks": ["CIS 9", "PCI DSS Req.1"],
                "remediation": "Restrict SMB to segmented admin networks",
            })
        if "windows" in os_name and 5985 in ports:
            findings.append({
                "rule": "Unrestricted WinRM HTTP",
                "severity": "medium",
                "frameworks": ["CIS Windows Benchmark"],
                "remediation": "Use HTTPS WinRM and limit remote hosts",
            })
        if not details.get("ports"):
            findings.append({
                "rule": "No Service Fingerprint Data",
                "severity": "low",
                "frameworks": ["Audit Coverage"],
                "remediation": "Run deep scan/service discovery",
            })
        return findings

    def generate_hardening_score(self, findings):
        score = 100
        penalty = {"critical": 30, "high": 20, "medium": 10, "low": 5}
        for finding in findings:
            score -= penalty.get(finding["severity"], 5)
        return max(score, 0)

    def remediate_findings(self, ip, findings):
        actions = []
        for finding in findings:
            if "Asset Naming Standard" in finding["rule"]:
                with self.state_lock:
                    if ip in self.devices and not self.devices[ip].get("custom_name"):
                        self.devices[ip]["custom_name"] = f"Asset-{ip.split('.')[-1]}"
                        self._save_to_db(ip, self.devices[ip])
                        actions.append("Auto-tagged unnamed asset")
            if "No Service Fingerprint Data" in finding["rule"]:
                actions.append("Queued deep scan recommendation")
        return actions

    def _run_config_assessment(self):
        self.log("🧭 Running configuration assessment...")
        self._queue_ui(self._start_task_ui, "Assessing CIS baseline and compliance controls...")
        with self.state_lock:
            device_snapshot = dict(self.devices)
            detail_snapshot = dict(self.device_details)

        lines = [
            f"🧭 CONFIGURATION ASSESSMENT - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            "Templates: CIS Baseline | PCI DSS | HIPAA | GDPR",
            "═══════════════════════════════════════════════════",
        ]
        results = {}
        for ip, info in sorted(device_snapshot.items()):
            details = detail_snapshot.get(ip, {})
            findings = self.check_compliance_rules(ip, info, details)
            score = self.generate_hardening_score(findings)
            results[ip] = {"score": score, "findings": findings}
            lines.append(f"\n{ip} ({info.get('custom_name', 'Unlabeled')}) - Hardening Score: {score}/100")
            if findings:
                for finding in findings[:6]:
                    lines.append(
                        f"  - [{finding['severity'].upper()}] {finding['rule']} | "
                        f"{'/'.join(finding['frameworks'])}"
                    )
                    lines.append(f"    Fix: {finding['remediation']}")
                worst = max(findings, key=lambda x: {"critical": 4, "high": 3, "medium": 2, "low": 1}.get(x["severity"], 1))
                level = {"critical": 16, "high": 13, "medium": 8, "low": 4}.get(worst["severity"], 4)
                self._emit_security_alert(level, f"Configuration drift detected on {ip}: {worst['rule']}", ip)
                if self.auto_remediate.get():
                    actions = self.remediate_findings(ip, findings)
                    if actions:
                        lines.append(f"    Auto-remediation: {', '.join(actions)}")
            else:
                lines.append("  - PASS: no baseline issues detected")

        self.config_assessment_results = results
        self._queue_ui(self._set_text_widget, self.endpoint_text, "\n".join(lines))
        self._queue_ui(self._end_task_ui)
        self.log("✅ Configuration assessment complete")

    def baseline_filesystem_gui(self):
        selected = filedialog.askdirectory(title="Select Directory To Baseline")
        target = selected or self.default_fim_path.get()
        threading.Thread(target=lambda: self._baseline_filesystem([target]), daemon=True).start()

    def _baseline_filesystem(self, paths):
        self.log(f"🧬 Creating filesystem baseline for: {', '.join(paths)}")
        self._queue_ui(self._start_task_ui, "Creating FIM baseline...")
        baseline = self.file_monitor.baseline_filesystem(paths)
        out = [
            f"🧬 FIM BASELINE - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            f"Paths: {', '.join(paths)}",
            f"Tracked files: {len(baseline)}",
            "Baseline ready. Start FIM monitor for continuous detection.",
        ]
        self._queue_ui(self._set_text_widget, self.fim_text, "\n".join(out))
        self._queue_ui(self._end_task_ui)

    def start_fim_monitor_gui(self):
        if not self.file_monitor.baseline:
            messagebox.showwarning("FIM", "Create a baseline first.")
            return
        if self.fim_running:
            messagebox.showinfo("FIM", "FIM monitor is already running.")
            return
        self.fim_running = True
        threading.Thread(target=self._run_fim_monitor, daemon=True).start()

    def stop_fim_monitor(self):
        self.fim_running = False
        self.log("🛑 FIM monitor stopped")

    def _run_fim_monitor(self):
        self.log("👁️ FIM monitor started")
        self._queue_ui(self._start_task_ui, "Monitoring filesystem integrity...")
        while self.fim_running:
            if self.monitoring_paused:
                time.sleep(1)
                continue
            changes = self.file_monitor.detect_file_changes()
            if changes:
                for ch in changes[:100]:
                    msg = f"{ch['type'].upper():<18} {ch['path']}"
                    self._queue_ui(self._append_text_widget, self.fim_text, msg + "\n")
                    self._emit_security_alert(ch["rule_level"], f"FIM {ch['type']} detected", ch["path"])
            time.sleep(max(5, int(self.fim_interval.get())))
        self._queue_ui(self._end_task_ui)

    def restore_from_backup(self, path):
        return {"path": path, "status": "manual_restore_required", "note": "Integrate with enterprise backup tooling"}

    def hygiene_scan_gui(self):
        threading.Thread(target=self._run_hygiene_scan, daemon=True).start()

    def hygiene_scanner(self, ip, info, details):
        findings = []
        ports = set(int(p) for p in details.get("ports", {}).keys())
        if ports & {21, 23}:
            findings.append("Legacy plaintext protocols enabled")
        if 80 in ports and 443 not in ports:
            findings.append("HTTP exposed without HTTPS")
        if 3389 in ports and not info.get("is_critical", False):
            findings.append("RDP exposed on non-critical asset")
        if info.get("hostname", "Unknown") == "Unknown":
            findings.append("Hostname unresolved")
        if not info.get("custom_name"):
            findings.append("No owner/label assigned")
        return findings

    def hygiene_scorecard(self, findings):
        return max(100 - len(findings) * 15, 0)

    def remediation_workflow(self, findings):
        priorities = []
        for finding in findings:
            if "plaintext" in finding or "RDP" in finding:
                priorities.append(f"[P1] {finding}")
            elif "HTTPS" in finding:
                priorities.append(f"[P2] {finding}")
            else:
                priorities.append(f"[P3] {finding}")
        return priorities

    def _run_hygiene_scan(self):
        self.log("🧹 Running IT hygiene assessment...")
        with self.state_lock:
            device_snapshot = dict(self.devices)
            detail_snapshot = dict(self.device_details)
        lines = ["🧹 IT HYGIENE SCORECARD", "═══════════════════════════════════════"]
        summary_score = []
        for ip, info in sorted(device_snapshot.items()):
            findings = self.hygiene_scanner(ip, info, detail_snapshot.get(ip, {}))
            score = self.hygiene_scorecard(findings)
            summary_score.append(score)
            lines.append(f"\n{ip} -> Hygiene Score: {score}/100")
            if findings:
                for item in self.remediation_workflow(findings):
                    lines.append(f"  {item}")
            else:
                lines.append("  ✅ No hygiene issues detected")
        overall = int(sum(summary_score) / max(len(summary_score), 1))
        lines.insert(1, f"Overall Hygiene Score: {overall}/100")
        self.hygiene_results = {"overall": overall}
        self._queue_ui(self._set_text_widget, self.endpoint_text, "\n".join(lines))

    def pci_scope_identifier(self):
        with self.state_lock:
            return [ip for ip, info in self.devices.items() if info.get("is_critical", False)]

    def pci_control_validator(self):
        scoped = self.pci_scope_identifier()
        controls = {"req1_network_segmentation": bool(scoped), "req10_logging": len(self.alert_manager.alerts) > 0}
        return controls

    def pci_report_generator(self):
        controls = self.pci_control_validator()
        passed = sum(1 for v in controls.values() if v)
        return {"framework": "PCI DSS", "passed_controls": passed, "total_controls": len(controls)}

    def personal_data_locator(self):
        patterns = [r"\b\d{3}-\d{2}-\d{4}\b", r"\b\d{16}\b", r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b"]
        findings = []
        with self.state_lock:
            for ip, events in self.traffic_logs.items():
                blob = "\n".join(events[-50:])
                for p in patterns:
                    if re.search(p, blob, re.IGNORECASE):
                        findings.append(ip)
                        break
        return sorted(set(findings))

    def consent_tracker(self):
        return {"status": "manual_process_required", "note": "Track consent records in governance workflow"}

    def data_breach_notification(self, impacted_records):
        if impacted_records >= 500:
            self._emit_security_alert(16, f"Potential reportable breach: {impacted_records} records impacted", "compliance")
        return impacted_records >= 500

    def right_to_be_forgotten(self, identifier):
        return {"identifier": identifier, "status": "queued_for_erasure_review"}

    def phi_locator(self):
        terms = ("diagnosis", "patient", "prescription", "medical")
        hits = []
        with self.state_lock:
            for ip, events in self.traffic_logs.items():
                joined = " ".join(events[-50:]).lower()
                if any(term in joined for term in terms):
                    hits.append(ip)
        return hits

    def hipaa_audit_controls(self):
        return {"audit_logging": True, "access_review": True, "phi_monitoring": bool(self.phi_locator())}

    def breach_notification_rule(self, record_count):
        return self.data_breach_notification(record_count)

    def business_associate_agreement_tracker(self):
        return {"status": "tracked_outside_tool", "last_review": datetime.now().strftime("%Y-%m-%d")}

    def threat_hunt_gui(self):
        threading.Thread(target=self._run_threat_hunt, daemon=True).start()

    def vulnerability_scanner(self, ip, details):
        findings = []
        for port in set(int(p) for p in details.get("ports", {}).keys()):
            if port in {21, 23, 445, 3389}:
                findings.append({"port": port, "cvss": 8.5, "exploit_available": True})
            elif port in {80, 161, 514}:
                findings.append({"port": port, "cvss": 6.8, "exploit_available": False})
        return findings

    def prioritization_engine(self, vuln_items, critical):
        prioritized = []
        for item in vuln_items:
            score = item["cvss"] * (1.4 if item["exploit_available"] else 1.0) * (1.3 if critical else 1.0)
            prioritized.append((score, item))
        return sorted(prioritized, key=lambda x: x[0], reverse=True)

    def _run_threat_hunt(self):
        self.log("🎯 Running threat hunt...")
        self._queue_ui(self._start_task_ui, "Hunting IOCs and mapping MITRE ATT&CK...")
        with self.state_lock:
            device_snapshot = dict(self.devices)
            detail_snapshot = dict(self.device_details)
            traffic_snapshot = dict(self.traffic_logs)
        lines = [
            f"🎯 THREAT HUNT REPORT - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            "IOC Hunt + MITRE ATT&CK Mapping",
            "═══════════════════════════════════════",
            self.trend_analysis(),
            self.sli_slo_tracking(),
            "",
            "Correlation Engine:",
        ]
        for c in self.alert_correlation_engine():
            lines.append(f"  - {c}")
        lines.append("")
        lines.append(self.alert_timeline_view())
        lines.append("\nDevice Threat Scores:")
        for ip, info in sorted(device_snapshot.items()):
            details = detail_snapshot.get(ip, {})
            score = self.threat_hunter.threat_score_device(details, info)
            technique_id, technique_name = self.threat_hunter.map_to_mitre(" ".join(traffic_snapshot.get(ip, [])[-5:]))
            vulns = self.vulnerability_scanner(ip, details)
            prioritized = self.prioritization_engine(vulns, info.get("is_critical", False))
            lines.append(f"\n{ip} -> Threat Score: {score}/100 | MITRE: {technique_id} {technique_name}")
            if prioritized:
                top = prioritized[0][1]
                lines.append(f"  Top vuln candidate: Port {top['port']} (CVSS {top['cvss']})")
        cloud_status = []
        for name, monitor in self.cloud_monitors.items():
            if name == "docker":
                cloud_status.append(f"  Docker: {monitor.monitor_containers()['status']}")
            elif name == "aws":
                cloud_status.append(f"  AWS: {monitor.monitor_aws_events()['status']}")
            else:
                cloud_status.append(f"  GitHub: {monitor.monitor_audit_log()['status']}")
        lines.append("\nCloud Security Integration Status:")
        lines.extend(cloud_status)
        self._queue_ui(self._set_text_widget, self.threat_text, "\n".join(lines))
        self._queue_ui(self._end_task_ui)
    
    def _update_services_tab(self):
        services_summary = ["🛠️ GLOBAL NETWORK SERVICES INVENTORY", "═══════════════════════════════════════"]
        for ip, details in self.device_details.items():
            ports = details.get('ports', {})
            if ports:
                name = self.devices.get(ip, {}).get('custom_name', 'No Name')
                services_summary.append(f"\nDevice: {ip} ({name})")
                for port, svc in ports.items():
                    services_summary.append(f"  • Port {port}: {svc}")
        
        self._set_text_widget(self.services_text, "\n".join(services_summary))

    def _deep_scan_device(self, ip):
        if not self._is_valid_ipv4(ip):
            self.log(f"⚠️ Invalid target IP skipped: {ip}", "WARN")
            return
        self.log(f"🔍 Deep scanning {ip}...")
        self._queue_ui(self._start_task_ui, f"Deep scanning {ip}...")
        details = self.device_details.get(ip, {})
        details['os'] = self._detect_os(ip)
        details['ports'] = self._quick_port_scan(ip)
        try:
            if platform.system() == 'Windows':
                cmd = ['tracert', '-d', '-h', '5', ip]
            else:
                cmd = ['traceroute', '-n', '-m', '5', ip]
            res = self._run_command(cmd, capture_output=True, timeout=12)
            details['traceroute'] = res.stdout if res else "Traceroute unavailable"
        except: details['traceroute'] = "Failed"
        self.device_details[ip] = details
        self._queue_ui(self._show_device_details, ip)
        self._queue_ui(self._update_services_tab)
        self._queue_ui(self._end_task_ui)
        self.log(f"✅ Deep scan complete: {ip}")

    def _detect_os(self, ip):
        if NMAP_AVAILABLE:
            try:
                nm = nmap.PortScanner()
                nm.scan(ip, arguments='-O --osscan-limit')
                if 'osmatch' in nm[ip] and nm[ip]['osmatch']:
                    return nm[ip]['osmatch'][0]['name']
            except: pass
        return "Unknown (Requires sudo/Nmap)"

    def _quick_port_scan(self, ip):
        common = [21,22,23,80,443,3389,8080]
        open_p = {}
        for port in common:
            with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as sock:
                sock.settimeout(0.3)
                if sock.connect_ex((ip, port)) == 0:
                    open_p[port] = self._get_service_name(port)
        return open_p

    def _get_service_name(self, port):
        services = {21:'FTP', 22:'SSH', 23:'Telnet', 80:'HTTP', 443:'HTTPS', 3389:'RDP'}
        return services.get(port, f'TCP/{port}')

    def nmap_scan_selected_gui(self):
        if not self.selected_device:
            messagebox.showwarning("No Selection", "Select a device first!")
            return
        if not self._is_valid_ipv4(self.selected_device):
            messagebox.showerror("Invalid IP", f"Cannot scan invalid IP: {self.selected_device}")
            return
        if not NMAP_AVAILABLE:
            messagebox.showerror("Missing Dependency", "python-nmap not found. Install nmap + python3-nmap.")
            return
        threading.Thread(target=lambda: self._nmap_full_scan_single(self.selected_device), daemon=True).start()

    def port_scan_selected(self):
        # Backward-compatible alias for older button/action wiring.
        self.nmap_scan_selected_gui()

    def _nmap_full_scan_single(self, ip):
        if not NMAP_AVAILABLE: return
        if not self._is_valid_ipv4(ip):
            self.log(f"⚠️ Invalid target IP skipped: {ip}", "WARN")
            return
        self.log(f"🌐 Nmap scanning {ip}...")
        self._queue_ui(self._start_task_ui, f"Nmap scan in progress ({ip})...")
        nm = nmap.PortScanner()
        nm.scan(ip, arguments='-sV -T4')
        res = [f"🌐 NMAP RESULTS: {ip}"]
        if ip in nm.all_hosts():
            details = self.device_details.get(ip, {})
            details['ports'] = details.get('ports', {})
            for proto in nm[ip].all_protocols():
                for port in nm[ip][proto]:
                    svc = f"{nm[ip][proto][port]['name']} ({nm[ip][proto][port]['version']})"
                    res.append(f"Port {port}: {svc}")
                    details['ports'][port] = svc
            self.device_details[ip] = details
        self._queue_ui(self._set_text_widget, self.results_text, "\n".join(res))
        self._queue_ui(self._show_device_details, ip)
        self._queue_ui(self._update_services_tab)
        self._queue_ui(self._end_task_ui)

    def traceroute_selected(self):
        if not self.selected_device: return
        if not self._is_valid_ipv4(self.selected_device):
            messagebox.showerror("Invalid IP", f"Cannot trace invalid IP: {self.selected_device}")
            return
        threading.Thread(target=lambda: self._traceroute_device(self.selected_device), daemon=True).start()
    
    def _traceroute_device(self, ip):
        if not self._is_valid_ipv4(ip):
            self.log(f"⚠️ Invalid traceroute target skipped: {ip}", "WARN")
            return
        self.log(f"📡 Traceroute {ip}...")
        self._queue_ui(self._start_task_ui, f"Running traceroute for {ip}...")
        try:
            cmd = ['tracert', '-d', ip] if platform.system() == 'Windows' else ['traceroute', '-n', ip]
            res = self._run_command(cmd, capture_output=True, timeout=15)
            if res is None:
                raise RuntimeError("Traceroute command unavailable")
            details = self.device_details.get(ip, {})
            details["traceroute"] = res.stdout
            self.device_details[ip] = details
            self._queue_ui(self._set_text_widget, self.results_text, res.stdout)
            self._queue_ui(self._show_device_details, ip)
        except Exception as exc:
            self.log(f"⚠️ Traceroute failed for {ip}: {exc}", "WARN")
        finally:
            self._queue_ui(self._end_task_ui)

    def nmap_scan_gui(self):
        if not NMAP_AVAILABLE or nmap is None:
            messagebox.showerror("Missing Dependency", "python-nmap not found. Install nmap + python3-nmap.")
            return
        threading.Thread(target=self._nmap_full_scan, daemon=True).start()

    def _nmap_full_scan(self):
        self.log("🌐 Scanning network range with Nmap...")
        self._queue_ui(self._start_task_ui, "Running network-wide Nmap scan...")
        try:
            nm = nmap.PortScanner()
            target = f"{self.router_ip}/24" if self.router_ip and self._is_valid_ipv4(self.router_ip) else "192.168.1.0/24"
            nm.scan(hosts=target, arguments='-F')
            
            discovered = 0
            for host in nm.all_hosts():
                if nm[host].state() == 'up':
                    discovered += 1
                    with self.state_lock:
                        info = self.devices.get(host, {
                            'mac': '??:??:??:??:??:??', 
                            'hostname': nm[host].hostname() or 'Unknown',
                            'vendor': 'Unknown', 
                            'status': 'LIVE',
                            'connection_time': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                        })
                        info['status'] = 'LIVE'
                        self.devices[host] = info
                    self._save_to_db(host, info)
            
            self._queue_ui(self._update_devices_tree)
            self.log(f"✅ Nmap range scan complete: found {discovered} LIVE hosts", "SUCCESS")
        except Exception as e:
            self.log(f"❌ Nmap range scan failed: {e}", "ERROR")
        finally:
            self._queue_ui(self._end_task_ui)

    def service_discovery_all_gui(self):
        threading.Thread(target=self._service_discovery_all, daemon=True).start()

    def _service_discovery_all(self):
        self.log("🛠️ Starting Global Service Discovery...")
        with self.state_lock:
            targets = [ip for ip, info in self.devices.items() if "LIVE" in info.get('status', '')]
        for ip in targets:
            self._deep_scan_device(ip)
        self.log("✅ Global Service Discovery complete")

    def bandwidth_monitor_gui(self):
        threading.Thread(target=self._bandwidth_monitor, daemon=True).start()
    
    def _bandwidth_monitor(self):
        self.log("📈 10s bandwidth test...")
        self._queue_ui(self._start_task_ui, "Measuring bandwidth (10s)...")
        s = psutil.net_io_counters()
        time.sleep(10)
        e = psutil.net_io_counters()
        up = (e.bytes_sent - s.bytes_sent) / 1024 / 1024 * 8 / 10
        down = (e.bytes_recv - s.bytes_recv) / 1024 / 1024 * 8 / 10
        result = f"BANDWIDTH RESULTS:\nUpload: {up:.2f} Mbps\nDownload: {down:.2f} Mbps"
        self._queue_ui(self._set_text_widget, self.results_text, result)
        self._queue_ui(self._end_task_ui)

    def packet_sniffer_gui(self):
        if not SCAPY_AVAILABLE or sniff is None:
            messagebox.showerror("Missing Dependency", "Scapy not available. Install with: pip install scapy")
            return
        threading.Thread(target=self._packet_sniffer, daemon=True).start()
    
    def _packet_sniffer(self):
        self.log("🌐 Sniffing 20 packets...")
        sniff(count=20, prn=lambda x: self.log(f"Packet: {x.summary()}"), store=False)

    def _start_traffic_monitor(self):
        if SCAPY_AVAILABLE and sniff is not None:
            threading.Thread(target=self._traffic_sniffer, daemon=True).start()
        else:
            self.log("⚠️ Activity monitor disabled: Scapy not available", "WARN")

    def _traffic_sniffer(self):
        if not SCAPY_AVAILABLE or sniff is None:
            return
        
        if not self._is_root_user():
            self.log("⚠️ ROOT REQUIRED for Traffic Monitoring!", "WARN")
            self._queue_ui(self._log_activity, "⚠️ Permission Denied: Run as sudo to monitor IP activity.")
            return

        def process_pkt(pkt):
            try:
                if self.monitoring_paused:
                    return
                if not pkt.haslayer(IP): return
                ip = pkt[IP].src
                
                # DNS Monitoring
                if pkt.haslayer(DNSQR):
                    query = pkt[DNSQR].qname.decode('utf-8', errors='ignore').rstrip('.')
                    query = self._sanitize_host(query)
                    with self.state_lock:
                        if ip not in self.traffic_logs:
                            self.traffic_logs[ip] = []
                    entry = f"{datetime.now().strftime('%H:%M:%S')} - DNS: {query}"
                    with self.state_lock:
                        if entry not in self.traffic_logs[ip]:
                            self.traffic_logs[ip].append(entry)
                            self._queue_ui(self._log_activity, f"🌐 {ip} -> {query}")
                
                # HTTP Monitoring (Port 80)
                elif pkt.haslayer(TCP) and pkt.haslayer(Raw):
                    payload = pkt[Raw].load.decode('utf-8', errors='ignore')
                    if "Host:" in payload:
                        host = re.search(r'Host: (.*?)\r\n', payload)
                        if host:
                            hostname = self._sanitize_host(host.group(1))
                            with self.state_lock:
                                if ip not in self.traffic_logs:
                                    self.traffic_logs[ip] = []
                            entry = f"{datetime.now().strftime('%H:%M:%S')} - HTTP: {hostname}"
                            with self.state_lock:
                                if entry not in self.traffic_logs[ip]:
                                    self.traffic_logs[ip].append(entry)
                                    self._queue_ui(self._log_activity, f"🌐 {ip} -> http://{hostname}")
            except: pass

        try:
            self.log("📡 Activity Monitor: STARTED (DNS/HTTP)")
            sniff(prn=process_pkt, store=False, filter="udp port 53 or tcp port 80")
        except Exception as e:
            self.log(f"❌ Sniffing Error: {e}", "ERROR")
            self._queue_ui(self._log_activity, f"❌ Error: {e}")

    def _log_activity(self, message):
        if threading.get_ident() != self.main_thread_id:
            self._queue_ui(self._log_activity, message)
            return
        timestamp = datetime.now().strftime("%H:%M:%S")
        entry = f"[{timestamp}] {message}\n"
        self.activity_text.insert(tk.END, entry)
        self.activity_text.see(tk.END)

    def save_all_data_gui(self):
        directory = filedialog.askdirectory(title="Select Directory to Save All Reports")
        if not directory:
            return
        
        self.log(f"📦 Exporting all reports to {directory}...")
        try:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            export_dir = os.path.join(directory, f"netmon_full_report_{timestamp}")
            os.makedirs(export_dir, exist_ok=True)
            
            # 1. Main Report (Text)
            report = self.generate_report()
            with open(os.path.join(export_dir, "summary_report.txt"), "w", encoding="utf-8") as f:
                f.write(report)
            
            # 2. Devices Inventory (CSV)
            with open(os.path.join(export_dir, "device_inventory.csv"), "w", encoding="utf-8", newline="") as f:
                writer = csv.writer(f)
                writer.writerow(['IP', 'MAC', 'Hostname', 'Vendor', 'Status', 'Custom Name', 'Is Critical'])
                for ip, info in self.devices.items():
                    writer.writerow([ip, info.get('mac'), info.get('hostname'), info.get('vendor'), 
                                     info.get('status'), info.get('custom_name', ''), info.get('is_critical', False)])
            
            # 3. Operations Log
            with open(os.path.join(export_dir, "operations.log"), "w", encoding="utf-8") as f:
                f.write(self.log_text.get(1.0, tk.END))
            
            # 4. Traffic Logs (JSON)
            with open(os.path.join(export_dir, "traffic_logs.json"), "w", encoding="utf-8") as f:
                json.dump(self.traffic_logs, f, indent=2)
            
            # 5. Full Device Details (JSON)
            combined_details = {}
            for ip in self.devices:
                combined_details[ip] = {
                    "info": self.devices.get(ip),
                    "details": self.device_details.get(ip)
                }
            with open(os.path.join(export_dir, "full_device_details.json"), "w", encoding="utf-8") as f:
                json.dump(combined_details, f, indent=2)

            self.log(f"✅ Full export complete in {export_dir}")
            messagebox.showinfo("Success", f"All data and reports exported to:\n{export_dir}")
        except Exception as exc:
            self.log(f"❌ Full export failed: {exc}", "ERROR")
            messagebox.showerror("Export Error", f"Failed to export all data: {exc}")

    def _safe_pdf_text(self, text):
        if text is None:
            return ""
        return str(text).encode("latin-1", "replace").decode("latin-1")

    def _build_pdf_report_lines(self):
        with self.state_lock:
            total = len(self.devices)
            live = sum(1 for info in self.devices.values() if "LIVE" in info.get("status", ""))
            crit = sum(1 for info in self.devices.values() if info.get("is_critical", False))
            devices_snapshot = dict(self.devices)
            details_snapshot = dict(self.device_details)
            traffic_snapshot = dict(self.traffic_logs)

        operations_log = self.log_text.get("1.0", tk.END).strip().splitlines() if hasattr(self, "log_text") else []
        activity_log = self.activity_text.get("1.0", tk.END).strip().splitlines() if hasattr(self, "activity_text") else []
        alert_count = len(getattr(self.alert_manager, "alerts", []))

        module_snapshots = [
            ("Scan Results", self.results_text.get("1.0", tk.END) if hasattr(self, "results_text") else ""),
            ("Security Audit", self.audit_text.get("1.0", tk.END) if hasattr(self, "audit_text") else ""),
            ("Endpoint Security", self.endpoint_text.get("1.0", tk.END) if hasattr(self, "endpoint_text") else ""),
            ("File Integrity", self.fim_text.get("1.0", tk.END) if hasattr(self, "fim_text") else ""),
            ("Threat Intel", self.threat_text.get("1.0", tk.END) if hasattr(self, "threat_text") else ""),
            ("Network Services", self.services_text.get("1.0", tk.END) if hasattr(self, "services_text") else ""),
            ("GRC Compliance", self.grc_text.get("1.0", tk.END) if hasattr(self, "grc_text") else ""),
        ]

        lines = [
            "NETMON PRO v2.1",
            "Enterprise Network Security Audit",
            "",
            f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            f"Total Assets Tracked: {total}",
            f"Live Assets Detected: {live}",
            f"Critical Assets Flagged: {crit}",
            f"GRC Compliance Score: {self.grc_results.get('score', 'N/A')}/100",
            f"Overall Risk Level: {self.grc_results.get('risk_level', 'N/A')}",
            "",
            "Operational Activity Summary",
            "",
            f"Total Operations Log Entries: {len(operations_log)}",
            f"Total Activity Monitor Events: {len(activity_log)}",
            f"Total Security Alerts Generated: {alert_count}",
            "",
            "Recent Operations (last 60)",
        ]

        for row in operations_log[-60:]:
            lines.append(f"  - {row}")

        lines.extend([
            "",
            "Recent Activity Monitor Events (last 60)",
        ])

        for row in activity_log[-60:]:
            lines.append(f"  - {row}")

        lines.extend([
            "",
            "Module Output Summary",
            "Asset Intelligence Detail",
            "",
        ])

        for title, content in module_snapshots:
            text = content.strip()
            if not text:
                continue
            snapshot_lines = text.splitlines()
            lines.append(f"{title}: {len(snapshot_lines)} lines")
            for row in snapshot_lines[-20:]:
                lines.append(f"  - {row}")
            lines.append("")

        for ip, info in devices_snapshot.items():
            name = info.get("custom_name", "Unnamed")
            crit_tag = "[CRITICAL] " if info.get("is_critical") else ""
            lines.append(f"{crit_tag}Asset: {ip} ({name})")
            lines.append(f"MAC: {info.get('mac', 'Unknown')} | Status: {info.get('status', 'Unknown')}")
            lines.append(f"Hostname: {info.get('hostname', 'Unknown')} | Vendor: {info.get('vendor', 'Unknown')}")
            lines.append(f"First Discovery: {info.get('connection_time', 'N/A')}")

            details = details_snapshot.get(ip, {})
            lines.append(f"OS Fingerprint: {details.get('os', 'Unknown')}")

            ports = details.get("ports", {})
            if ports:
                lines.append("Open Services:")
                for port, svc in ports.items():
                    lines.append(f"  - Port {port}: {svc}")

            ip_logs = traffic_snapshot.get(ip, [])
            if ip_logs:
                lines.append("Security Event Log (Traffic):")
                for log in ip_logs[-10:]:
                    lines.append(f"  > {log}")
            lines.append("")
        return lines

    def _write_basic_pdf(self, path, lines):
        width = 595
        height = 842
        left_margin = 40
        top_margin = 42
        line_height = 14
        max_lines = max(1, int((height - (top_margin * 2)) / line_height))

        sanitized_lines = [self._safe_pdf_text(line) for line in lines]
        pages = [sanitized_lines[i:i + max_lines] for i in range(0, len(sanitized_lines), max_lines)] or [[""]]

        object_defs = {}
        object_streams = {}
        next_obj_id = 1

        catalog_id = next_obj_id
        next_obj_id += 1
        pages_id = next_obj_id
        next_obj_id += 1
        font_id = next_obj_id
        next_obj_id += 1

        page_ids = []
        for page_lines in pages:
            page_id = next_obj_id
            next_obj_id += 1
            content_id = next_obj_id
            next_obj_id += 1
            page_ids.append(page_id)

            content_cmds = ["BT", "/F1 11 Tf", f"1 0 0 1 {left_margin} {height - top_margin} Tm"]
            first_line = True
            for line in page_lines:
                escaped = line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
                if not first_line:
                    content_cmds.append(f"0 -{line_height} Td")
                content_cmds.append(f"({escaped}) Tj")
                first_line = False
            content_cmds.append("ET")
            stream_bytes = "\n".join(content_cmds).encode("latin-1", "replace")
            object_streams[content_id] = stream_bytes
            object_defs[content_id] = f"<< /Length {len(stream_bytes)} >>"

            object_defs[page_id] = (
                f"<< /Type /Page /Parent {pages_id} 0 R "
                f"/MediaBox [0 0 {width} {height}] "
                f"/Resources << /Font << /F1 {font_id} 0 R >> >> "
                f"/Contents {content_id} 0 R >>"
            )

        kids = " ".join(f"{pid} 0 R" for pid in page_ids)
        object_defs[catalog_id] = f"<< /Type /Catalog /Pages {pages_id} 0 R >>"
        object_defs[pages_id] = f"<< /Type /Pages /Kids [{kids}] /Count {len(page_ids)} >>"
        object_defs[font_id] = "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>"

        max_id = next_obj_id - 1
        pdf = bytearray(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
        offsets = [0] * (max_id + 1)

        for obj_id in range(1, max_id + 1):
            offsets[obj_id] = len(pdf)
            pdf.extend(f"{obj_id} 0 obj\n".encode("ascii"))
            pdf.extend(object_defs[obj_id].encode("latin-1"))
            if obj_id in object_streams:
                pdf.extend(b"\nstream\n")
                pdf.extend(object_streams[obj_id])
                pdf.extend(b"\nendstream")
            pdf.extend(b"\nendobj\n")

        xref_offset = len(pdf)
        pdf.extend(f"xref\n0 {max_id + 1}\n".encode("ascii"))
        pdf.extend(b"0000000000 65535 f \n")
        for obj_id in range(1, max_id + 1):
            pdf.extend(f"{offsets[obj_id]:010d} 00000 n \n".encode("ascii"))

        pdf.extend(
            (
                f"trailer\n<< /Size {max_id + 1} /Root {catalog_id} 0 R >>\n"
                f"startxref\n{xref_offset}\n%%EOF\n"
            ).encode("ascii")
        )

        with open(path, "wb") as f:
            f.write(pdf)

    def generate_pdf_report(self):
        path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF files", "*.pdf")])
        if not path:
            return

        try:
            lines = self._build_pdf_report_lines()
            fpdf_class = _resolve_fpdf_class()

            if fpdf_class is not None:
                pdf = fpdf_class(orientation="P", unit="mm", format="A4")
                pdf.set_auto_page_break(auto=True, margin=15)
                pdf.add_page()
                pdf.set_font("Helvetica", size=10)

                for line in lines:
                    text = self._safe_pdf_text(line)
                    if text:
                        pdf.multi_cell(190, 6, text)
                    else:
                        pdf.ln(3)
                pdf.output(path)
            else:
                self._write_basic_pdf(path, lines)
                self.log("⚠️ fpdf/fpdf2 not installed; used built-in PDF exporter", "WARN")

            self.log(f"✅ PDF Audit Report saved: {path}", "SUCCESS")
            messagebox.showinfo("Success", f"PDF Audit Report saved to {path}")
        except Exception as e:
            self.log(f"❌ PDF Export Error: {e}", "ERROR")
            messagebox.showerror(
                "Export Error",
                f"Failed to save PDF: {e}\n\nOptional dependency: pip install fpdf2",
            )

    def generate_docx_report(self):
        if not DOCX_AVAILABLE:
            messagebox.showerror("Error", "python-docx library not found. Install with: pip install python-docx")
            return
        
        path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word documents", "*.docx")])
        if not path: return

        try:
            doc = docx.Document()
            doc.add_heading('🛡️ NETMON PRO v2.1 - Network Report', 0)
            doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'IP Address'
            hdr_cells[1].text = 'Name'
            hdr_cells[2].text = 'MAC Address'
            hdr_cells[3].text = 'Hostname'
            hdr_cells[4].text = 'Vendor'

            for ip, info in self.devices.items():
                row_cells = table.add_row().cells
                row_cells[0].text = ip
                row_cells[1].text = info.get('custom_name', '')
                row_cells[2].text = info.get('mac', '')
                row_cells[3].text = info.get('hostname', '')
                row_cells[4].text = info.get('vendor', '')

            doc.add_heading('Device Activity Log', level=1)
            for ip, logs in self.traffic_logs.items():
                if logs:
                    doc.add_heading(f"Activity for {ip} ({self.devices.get(ip, {}).get('custom_name', 'N/A')})", level=2)
                    for log in logs[-10:]:
                        doc.add_paragraph(log, style='List Bullet')

            doc.save(path)
            self.log(f"✅ DOCX Report saved: {path}")
        except Exception as e:
            self.log(f"❌ DOCX Export Error: {e}", "ERROR")

    def scan_processes_gui(self):
        threading.Thread(target=self._scan_processes, daemon=True).start()
    
    def _scan_processes(self):
        self.log("🐛 Scanning processes...")
        self._queue_ui(self._start_task_ui, "Scanning suspicious processes...")
        alerts = []
        for proc in psutil.process_iter(['pid', 'name', 'cmdline']):
            try:
                cmd = ' '.join(proc.info['cmdline'] or [])
                if any(s in cmd for s in ['nc', 'netcat', 'bash -i', 'nmap']):
                    alerts.append(f"🚨 {proc.info['name']} (PID:{proc.info['pid']})")
            except: pass
        self._queue_ui(self._set_text_widget, self.results_text, "\n".join(alerts) if alerts else "✅ Clean")
        self._queue_ui(self._end_task_ui)

    def generate_report(self):
        with self.state_lock:
            total_devices = len(self.devices)
            critical_assets = sum(1 for info in self.devices.values() if info.get("is_critical", False))
        grc_score = self.grc_results.get("score", "N/A")
        grc_risk = self.grc_results.get("risk_level", "N/A")
        hygiene_score = self.hygiene_results.get("overall", "N/A")
        alert_counts = self.alert_manager.counts_last_24h()
        report = (
            "🛡️ NETMON PRO v2.1 REPORT\n"
            f"Generated: {datetime.now()}\n"
            f"Devices: {total_devices}\n"
            f"Critical Assets: {critical_assets}\n"
            f"GRC Score: {grc_score}\n"
            f"GRC Risk: {grc_risk}\n"
            f"Hygiene Score: {hygiene_score}\n"
            f"Alerts (24h): Critical={alert_counts.get('critical',0)}, High={alert_counts.get('high',0)}, "
            f"Medium={alert_counts.get('medium',0)}, Low={alert_counts.get('low',0)}"
        )
        self._set_text_widget(self.results_text, report)
        return report

    def save_report_txt_gui(self):
        report = self.generate_report()
        path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("Text files", "*.txt"), ("All files", "*.*")],
        )
        if not path:
            return
        try:
            with open(path, "w", encoding="utf-8") as f:
                f.write(report + "\n")
            self.log(f"✅ Report saved: {path}")
            messagebox.showinfo("Success", f"Report saved to {path}")
        except Exception as exc:
            self.log(f"❌ Failed to save report: {exc}", "ERROR")
            messagebox.showerror("Save Error", f"Failed to save report: {exc}")

    def generate_framework_report_gui(self):
        framework = simpledialog.askstring(
            "Framework Report",
            "Enter framework (pci, hipaa, gdpr, iso27001/all):",
            initialvalue="pci",
        )
        if framework is None:
            return
        self._run_framework_report(framework)

    def export_csv(self):
        path = filedialog.asksaveasfilename(defaultextension=".csv")
        if path:
            try:
                with open(path, 'w', encoding='utf-8', newline='') as f:
                    writer = csv.writer(f)
                    writer.writerow(['IP', 'MAC', 'Hostname', 'Vendor'])
                    for ip, info in self.devices.items():
                        writer.writerow([ip, info.get('mac'), info.get('hostname'), info.get('vendor')])
                self.log(f"✅ CSV exported: {path}")
            except Exception as exc:
                self.log(f"❌ CSV export failed: {exc}", "ERROR")
                messagebox.showerror("Export Error", f"CSV export failed: {exc}")

    def save_device_details(self):
        if self.selected_device:
            path = filedialog.asksaveasfilename(defaultextension=".json")
            if path:
                try:
                    with open(path, 'w', encoding='utf-8') as f:
                        json.dump({'ip': self.selected_device, 'data': self.devices.get(self.selected_device)}, f, indent=2)
                    self.log(f"✅ Device details saved: {path}")
                except Exception as exc:
                    self.log(f"❌ JSON save failed: {exc}", "ERROR")
                    messagebox.showerror("Save Error", f"Device details save failed: {exc}")

    def toggle_critical_asset(self):
        if not self.selected_device: return
        ip = self.selected_device
        if ip in self.devices:
            self.devices[ip]['is_critical'] = not self.devices[ip].get('is_critical', False)
            self._save_to_db(ip, self.devices[ip])
            self._update_devices_tree()
            status = "MARKED AS CRITICAL" if self.devices[ip]['is_critical'] else "REMOVED FROM CRITICAL"
            self.log(f"🛡️ {ip} {status}")
            self._show_device_details(ip)

    def _detect_arp_anomalies(self):
        mac_to_ip = {}
        with self.state_lock:
            snapshot = list(self.devices.items())
        for ip, info in snapshot:
            mac = info.get("mac")
            if not mac or "??:" in mac:
                continue
            mac_to_ip.setdefault(mac, set()).add(ip)
        return sum(1 for ips in mac_to_ip.values() if len(ips) > 1)

    def _build_grc_telemetry(self):
        with self.state_lock:
            device_snapshot = dict(self.devices)
            detail_snapshot = dict(self.device_details)
            dns_events = sum(len(v) for v in self.traffic_logs.values())

        risky_ports = {21, 23, 80, 161, 445, 514, 3389}
        risky_findings = 0
        scanned_assets = 0
        for details in detail_snapshot.values():
            ports = details.get("ports", {})
            if ports:
                scanned_assets += 1
            for port in ports:
                if int(port) in risky_ports:
                    risky_findings += 1

        critical_assets = sum(1 for info in device_snapshot.values() if info.get("is_critical", False))
        named_assets = sum(1 for info in device_snapshot.values() if info.get("custom_name", "").strip())
        total_assets = len(device_snapshot)
        return {
            "total_assets": total_assets,
            "critical_assets": critical_assets,
            "named_assets": named_assets,
            "scanned_assets": scanned_assets,
            "risky_findings": risky_findings,
            "arp_anomalies": self._detect_arp_anomalies(),
            "dns_events": dns_events,
        }

    def _grc_control_status(self, telemetry):
        assets = telemetry["total_assets"] or 1
        name_ratio = telemetry["named_assets"] / assets
        scan_ratio = telemetry["scanned_assets"] / assets
        controls = {
            "ID.AM-1 / CIS 1.1 (Asset Inventory)": (
                "PASS" if name_ratio >= 0.8 else "WARN",
                f"{telemetry['named_assets']}/{telemetry['total_assets']} assets labeled",
            ),
            "PR.IP-3 / PCI DSS 1.2 (Insecure Service Exposure)": (
                "PASS" if telemetry["risky_findings"] == 0 else "FAIL",
                f"{telemetry['risky_findings']} risky port findings",
            ),
            "DE.CM-1 / CIS 13 (Network Monitoring)": (
                "PASS" if telemetry["dns_events"] > 0 else "WARN",
                f"{telemetry['dns_events']} DNS/HTTP activity events",
            ),
            "PR.AC-4 / ISO 27001 A.5.15 (Critical Asset Governance)": (
                "PASS" if telemetry["critical_assets"] > 0 else "WARN",
                f"{telemetry['critical_assets']} critical assets defined",
            ),
            "DE.CM-7 / NIST SI-4 (ARP/MITM Detection)": (
                "PASS" if telemetry["arp_anomalies"] == 0 else "FAIL",
                f"{telemetry['arp_anomalies']} ARP anomaly clusters",
            ),
            "PR.PT-1 / CIS 7 (Vulnerability Management Cadence)": (
                "PASS" if scan_ratio >= 0.6 else "WARN",
                f"{telemetry['scanned_assets']}/{telemetry['total_assets']} assets service-scanned",
            ),
        }
        return controls

    def _score_grc_controls(self, controls):
        weights = {"PASS": 100, "WARN": 60, "FAIL": 15}
        total = 0
        for status, _ in controls.values():
            total += weights.get(status, 0)
        score = int(total / max(len(controls), 1))
        if score >= 85:
            risk = "LOW"
        elif score >= 65:
            risk = "MEDIUM"
        else:
            risk = "HIGH"
        return score, risk

    def _refresh_grc_kpi(self, score=None, risk=None):
        score_text = f"GRC Score: {score}" if score is not None else "GRC Score: --"
        risk_text = f"Risk Level: {risk}" if risk else "Risk Level: --"
        risk_color = "#3fb950" if risk == "LOW" else "#d29922" if risk == "MEDIUM" else "#f85149"
        self.grc_score_label.config(text=score_text)
        self.grc_risk_label.config(text=risk_text, fg=risk_color if risk else "#ff7b72")

    def grc_audit_gui(self):
        threading.Thread(target=self._run_grc_audit, daemon=True).start()

    def _run_grc_audit(self):
        self.log("📘 Running GRC compliance assessment...")
        self._queue_ui(self._start_task_ui, "Running NIST/ISO/PCI/CIS compliance checks...")

        telemetry = self._build_grc_telemetry()
        controls = self._grc_control_status(telemetry)
        score, risk = self._score_grc_controls(controls)

        report_lines = [
            f"📘 GRC COMPLIANCE ASSESSMENT - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            "Framework Mapping: NIST CSF | ISO 27001 | PCI DSS | CIS Controls",
            "═══════════════════════════════════════════════════════════════════════",
            f"Overall Compliance Score: {score}/100",
            f"Residual Risk Level: {risk}",
            "",
            "Control Results:",
        ]

        for control_name, (status, detail) in controls.items():
            icon = "✅" if status == "PASS" else "⚠️" if status == "WARN" else "❌"
            report_lines.append(f"{icon} {control_name}: {status} ({detail})")

        report_lines.extend([
            "",
            "Audit Recommendations:",
            "1. Run deep scans on all production assets weekly.",
            "2. Remove/segment insecure services (Telnet/FTP/SMB/RDP) behind strict access controls.",
            "3. Enforce naming and ownership tags for every critical asset.",
            "4. Investigate any ARP anomaly immediately as potential MITM activity.",
            "5. Maintain continuous traffic monitoring and alert review workflows.",
        ])

        self.grc_results = {
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "score": score,
            "risk_level": risk,
            "telemetry": telemetry,
            "controls": controls,
        }
        self._queue_ui(self._set_text_widget, self.grc_text, "\n".join(report_lines))
        self._queue_ui(self._refresh_grc_kpi, score, risk)
        self._queue_ui(self._update_overview_dashboard)
        self._queue_ui(self._end_task_ui)
        self.log(f"✅ GRC assessment complete: score={score}, risk={risk}")

    def security_audit_gui(self):
        threading.Thread(target=self._run_security_audit, daemon=True).start()

    def _run_security_audit(self):
        self.log("🏦 Starting Bank-Grade Security Audit...")
        self._queue_ui(self._start_task_ui, "Running security audit...")
        findings = [f"🏦 SECURITY AUDIT REPORT - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", 
                    "═══════════════════════════════════════"]
        
        risky_ports = {
            21: "FTP (Insecure - Plaintext Credentials)",
            23: "Telnet (Critical Risk - Use SSH instead)",
            445: "SMB (Potential Ransomware Vector/EternalBlue)",
            3389: "RDP (Brute Force Target - Ensure MFA/VPN)",
            80: "HTTP (Insecure - Use HTTPS)",
            161: "SNMP (Potential Information Leak)",
            514: "Syslog (Potential Log Spoofing)"
        }

        found_risks = 0
        with self.state_lock:
            detail_snapshot = dict(self.device_details)
            device_snapshot = dict(self.devices)
        for ip, details in detail_snapshot.items():
            ports = details.get('ports', {})
            device_risks = []
            for port, svc in ports.items():
                if int(port) in risky_ports:
                    device_risks.append(f"  🚨 Port {port}: {risky_ports[int(port)]}")
            
            if device_risks:
                name = device_snapshot.get(ip, {}).get('custom_name', 'Unknown')
                findings.append(f"\nDevice: {ip} ({name})")
                findings.extend(device_risks)
                found_risks += 1
                self._emit_security_alert(13, "Insecure service exposure identified by security audit", ip)

        if found_risks == 0:
            findings.append("\n✅ No critical service risks identified on scanned devices.")
        else:
            findings.append(f"\n⚠️ Total {found_risks} devices with security concerns.")

        self.last_audit_summary = {
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "risky_devices": found_risks,
            "framework_hint": "Run GRC Compliance Audit for NIST/ISO/PCI/CIS mapping.",
        }
        findings.append("\n📘 Tip: Run 'GRC Compliance Audit' to map findings to NIST/ISO/PCI/CIS controls.")

        self._queue_ui(self._set_text_widget, self.audit_text, "\n".join(findings))
        self._queue_ui(self._end_task_ui)
        self.log(f"✅ Security Audit Complete: {found_risks} alerts")

    def asset_watchdog_gui(self):
        threading.Thread(target=self._run_asset_watchdog, daemon=True).start()

    def _run_asset_watchdog(self):
        with self.state_lock:
            critical_ips = [ip for ip, info in self.devices.items() if info.get('is_critical', False)]
        if not critical_ips:
            self._queue_ui(messagebox.showinfo, "Watchdog", "No critical assets defined. Mark a device as 'Critical' first.")
            return

        self.log(f"🛡️ Watchdog monitoring {len(critical_ips)} assets...")
        self._queue_ui(self._start_task_ui, "Monitoring critical assets...")
        results = ["🛡️ CRITICAL ASSET WATCHDOG STATUS", "═══════════════════════════════════════"]
        ping_cmd = ['ping', '-n', '1', '-w', '1000'] if platform.system() == 'Windows' else ['ping', '-c', '1', '-W', '1']
        
        for ip in critical_ips:
            try:
                if not self._is_valid_ipv4(ip):
                    results.append(f"⚪ {ip:<15} | INVALID | Skipped invalid IP")
                    continue
                res = self._run_command([*ping_cmd, ip], capture_output=True, timeout=2, silent=True)
                if res is None:
                    results.append(f"⚪ {ip:<15} | ERROR   | Ping command unavailable")
                    continue
                if res.returncode == 0:
                    ms = re.search(r'time=(.*?) ms', res.stdout)
                    latency = ms.group(1) if ms else "N/A"
                    results.append(f"🟢 {ip:<15} | ONLINE  | Latency: {latency}ms")
                else:
                    results.append(f"🔴 {ip:<15} | OFFLINE | ALERT: Core System Down!")
                    self.log(f"🚨 ALERT: Critical Asset {ip} is OFFLINE!", "ERROR")
                    self._emit_security_alert(15, "Critical asset offline", ip)
            except:
                results.append(f"⚪ {ip:<15} | ERROR   | Execution Failed")

        self._queue_ui(self._set_text_widget, self.results_text, "\n".join(results))
        self._queue_ui(self._end_task_ui)

    def arp_poison_detect_gui(self):
        threading.Thread(target=self._run_arp_poison_detect, daemon=True).start()

    def _run_arp_poison_detect(self):
        self.log("🕵️ Monitoring for ARP Spoofing/Poisoning...")
        self._queue_ui(self._start_task_ui, "Detecting ARP anomalies...")
        mac_to_ip = {}
        ip_to_mac = {}
        alerts = []

        # Check existing device list for anomalies
        with self.state_lock:
            snapshot = list(self.devices.items())
        for ip, info in snapshot:
            mac = info.get('mac')
            if not mac or "??:" in mac: continue
            
            if mac not in mac_to_ip: mac_to_ip[mac] = []
            if ip not in ip_to_mac: ip_to_mac[ip] = mac
            
            if ip not in mac_to_ip[mac]: mac_to_ip[mac].append(ip)

        for mac, ips in mac_to_ip.items():
            if len(ips) > 1:
                alerts.append(f"🚨 MITM ALERT: MAC {mac} is associated with multiple IPs: {', '.join(ips)}")
                for ip in ips:
                    self._emit_security_alert(16, "Potential ARP spoofing/MITM anomaly", ip)

        if alerts:
            self._queue_ui(self._set_text_widget, self.audit_text, "🕵️ ARP ANOMALY REPORT\n" + "\n".join(alerts))
            for a in alerts: self.log(a, "ERROR")
        else:
            self._queue_ui(self._set_text_widget, self.audit_text, "✅ No ARP anomalies detected in current device cache.")
        self._queue_ui(self._end_task_ui)

    def clear_all(self):
        if not messagebox.askyesno("Clear Data", "Clear in-memory data and delete device history from SQLite?"):
            return
        with self.state_lock:
            self.devices.clear()
            self.device_details.clear()
            self.traffic_logs.clear()
            self.grc_results.clear()
            self.last_audit_summary.clear()
            self.config_assessment_results.clear()
            self.hygiene_results.clear()
        self.fim_running = False
        self.alert_manager.alerts.clear()
        self.file_monitor.baseline.clear()
        self.file_monitor.monitored_paths.clear()
        try:
            with sqlite3.connect(DB_FILE) as conn:
                conn.execute("DELETE FROM devices")
        except Exception as exc:
            self.log(f"⚠️ Failed to clear database history: {exc}", "WARN")
        self.devices_tree.delete(*self.devices_tree.get_children())
        self._set_text_widget(self.results_text, "")
        self._set_text_widget(self.audit_text, "")
        self._set_text_widget(self.services_text, "")
        self._set_text_widget(self.activity_text, "")
        self._set_text_widget(self.grc_text, "")
        self._set_text_widget(self.alerts_text, "")
        self._set_text_widget(self.endpoint_text, "")
        self._set_text_widget(self.fim_text, "")
        self._set_text_widget(self.threat_text, "")
        self._refresh_grc_kpi()
        self._update_overview_dashboard()
        self.log("🧹 All data and history cleared")

def main():
    root = tk.Tk()
    app = NetMonPro(root)
    root.mainloop()

if __name__ == "__main__":
    main()
